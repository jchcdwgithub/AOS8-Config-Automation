from re import I, fullmatch
import AOS8_config_automation as CA
import setup
import pytest
from docx import Document

API_FILES_TEST = setup.get_API_JSON_files()

def test_build_hierarchy_builds_all_paths():
    """ The build hierarchy function correctly returns all intermediate paths. """

    expected = ['/md/NA','/md/NA/US','/md/NA/US/HQ', '/md/NA/US/HQ/DC']
    generated = CA.build_hierarchy('/NA/US/HQ/DC')

    assert expected == generated

def test_get_columns_from_tables_returns_all_columns():
    """ All the columns in the provided tables should be present in the returned columns dictionary. """

    tables = Document('../Radio Testing Tables.docx').tables

    expected = {"ap_g_radio_prof.profile-name":["CGR","DEP","KFD","PWK","SIG","Test"],
                "ap_a_radio_prof.profile-name":["CGR","DEP","KFD","PWK","SIG","Test"],
                "reg_domain_prof.profile-name":["CGR","DEP","KFD","PWK","SIG","Test"],
                "ap_group.profile-name":["CGR","DEP","KFD","PWK","SIG","Test"],
                "ap_g_radio_prof.eirp_min.eirp-min":["3 dBm", "3 dBm", "3 dBm", "3 dBm", "3 dBm", "3 dBm"],
                "ap_g_radio_prof.eirp_max.eirp-max":["7 dBm", "7 dBm", "7 dBm", "7 dBm", "7 dBm", "7 dBm"],
                "ap_a_radio_prof.eirp_min.eirp-min":["6 dBm", "6 dBm", "6 dBm", "6 dBm", "6 dBm", "6 dBm"],
                "ap_a_radio_prof.eirp_max.eirp-max":["10 dBm", "10 dBm", "10 dBm", "10 dBm", "10 dBm", "10 dBm"],
                "reg_domain_prof.valid_11b_channel.valid-11g-channel":["1, 6, 11", "1,6,11", "1, 6, 11", "1, 6, 11", "1, 6, 11", "1, 6, 11"],
                "reg_domain_prof.valid_11a_channel.valid-11a-channel":["36, 40, 44, 48, 52, 56, 60, 64, 149, 153, 157, 161,165", "36, 40, 44, 48, 52, 56, 60, 64, 149, 153, 157, 161,165", "36, 40, 44, 48, 52, 56, 60, 64, 149, 153, 157, 161,165", "36, 40, 44, 48, 52, 56, 60, 64, 149, 153, 157, 161,165", "36, 40, 44, 48, 52, 56, 60, 64, 149, 153, 157, 161,165", "36, 40, 44, 48, 149, 153, 157, 161,165" ],
                "ssid_prof.profile-name":["ChoiceAccess", "CorpAccess", "CorpTest", "GuestAccess", "MobiAccess"],
                "ssid_prof.essid.essid":["ChoiceAccess", "CorpAccess", "CorpTest", "GuestAccess", "MobiAccess"],
                "ssid_prof.g_basic_rates":["12", "12", "12", "12", "12"],
                "ssid_prof.g_tx_rates":["12, 18, 24, 36, 48, 54", "12, 18, 24, 36, 48, 54", "12, 18, 24, 36, 48, 54", "12, 18, 24, 36, 48, 54", "12, 18, 24, 36, 48, 54"],
                "ssid_prof.a_basic_rates":["12", "12", "12, 24", "12", "12"],
                "ssid_prof.a_tx_rates":["12, 18, 24, 36, 48, 54", "12, 18, 24, 36, 48, 54", "12, 18, 24, 36, 48, 54", "12, 18, 24, 36, 48, 54", "12, 18, 24, 36, 48, 54"]}
    
    CA.build_tables_columns_dict(tables)    
    CA.remove_column_headers_from_columns_table()

    assert expected == CA.TABLE_COLUMNS

def test_build_tables_columns_dict_builds_comma_separated_name_node_pairs():

    tables = Document('../Node Testing Tables.docx').tables

    expected = {"ap_g_radio_prof.profile-name":["Global,/md/HQ","HQ,/md/HQ","WAREHOUSE,/md/Branch1","FACTORY,/md/Branch2"],
                "ap_a_radio_prof.profile-name":["Global,/md/HQ","HQ,/md/HQ","WAREHOUSE,/md/Branch1","FACTORY,/md/Branch2"],
                "reg_domain_prof.profile-name":["Global,/md/HQ","HQ,/md/HQ","WAREHOUSE,/md/Branch1","FACTORY,/md/Branch2"],
                "ap_g_radio_prof.eirp_min.eirp-min":["8 dBm", "5 dBm", "11 dBm", "8 dBm"],
                "ap_g_radio_prof.eirp_max.eirp-max":["14 dBm", "11 dBm", "14 dBm", "14 dBm"],
                "ap_a_radio_prof.eirp_min.eirp-min":["11 dBm", "8 dBm", "11 dBm", "11 dBm"],
                "ap_a_radio_prof.eirp_max.eirp-max":["17 dBm", "14 dBm", "16 dBm", "14 dBm"]}
    
    CA.add_entries_to_object_identifiers()
    CA.build_tables_columns_dict(tables)
    CA.remove_column_headers_from_columns_table()

    assert expected == CA.TABLE_COLUMNS

def test_get_profiles_to_be_configured_returns_correct_dictionary():
    """ Given the above test tables, get the names and attributes of profiles to be configured. """

    tables = Document('../Radio Testing Tables.docx').tables
    CA.add_entries_to_object_identifiers()
    CA.build_tables_columns_dict(tables)
    CA.add_entries_to_object_identifiers()

    expected = {"ap_g_radio_prof":["profile-name","eirp_min.eirp-min","eirp_max.eirp-max"],
                "ap_a_radio_prof":["profile-name","eirp_min.eirp-min","eirp_max.eirp-max"],
                "ap_group":["profile-name"],
                "reg_domain_prof":["profile-name","valid_11b_channel.valid-11g-channel","valid_11a_channel.valid-11a-channel"],
                "ssid_prof":["profile-name","essid.essid","g_basic_rates","g_tx_rates","a_basic_rates","a_tx_rates"]}
    
    generated = CA.get_profiles_to_be_configured()

    assert expected == generated
    
def test_check_for_required_attributes_returns_true_when_all_attributes_are_configured():

    profile_name = "configuration_device_filename"
    profile_attributes = [
                "dev-model",
                "config-path",
                "mac-address",
                "filename"
            ]
    
    Result = CA.check_that_required_attributes_are_provided(profile_name,profile_attributes,API_FILES_TEST)

    assert Result == True

def test_check_for_required_attributes_returns_false_when_attributes_are_missing():

    profile_name = 'configuration_device_filename'
    profile_attributes = [
                "dev-model",
                "mac-address",
                "filename"
            ]

    Result = CA.check_that_required_attributes_are_provided(profile_name,profile_attributes,API_FILES_TEST)

    assert Result == False

def test_is_nested_profile_returns_true_for_internal_profile():

    attribute_name = "aaa_prof"

    Result = CA.is_nested_profile(attribute_name)

    assert Result == True

def test_is_nested_profile_returns_false_for_other_attribute_names():

    attribute_name = "eirp_max"

    Result = CA.is_nested_profile(attribute_name)

    assert Result == False

def test_get_attribute_type_returns_the_correct_integer_type():

    full_attribute_name = "foreign_agent_general_prof_it.interval-time"

    expected = "integer"

    generated = CA.get_attribute_type(full_attribute_name)

    assert expected == generated

def test_get_attribute_type_returns_the_correct_string_type():

    full_attribute_name = "ssid_prof.profile-name"

    expected = "string"

    generated = CA.get_attribute_type(full_attribute_name)

    assert expected == generated

def test_get_object_property_name_returns_a_list_of_names():

    full_attribute_name = "ap_mesh_radio_prof.mesh_a_tx_rates"

    expected = ["6", "9", "12", "18", "24", "36", "48", "54"]

    generated = CA.get_object_property_name(full_attribute_name)

    assert expected == generated

def test_is_enum_returns_true_for_enumerated_properties():

    profile_name = "anqp_nwk_auth_prof"
    attribute_name = "anqp_nwk_auth_type"
    property_name = CA.get_object_property_name(profile_name+'.'+attribute_name)[0]

    Result = CA.is_enumerated_property(f'{profile_name}.{attribute_name}.{property_name}')

    assert Result == True

def test_is_enum_returns_false_for_other_properties():

    full_attribute_name = "virtual_ap.aaa_prof.profile-name"

    Result = CA.is_enumerated_property(full_attribute_name)

    assert Result == False 

def test_string_is_in_enumerated_properties_list_returns_true_for_strings_in_list():

    profile_name = "anqp_nwk_auth_prof.anqp_nwk_auth_type.anqp_nwk_auth_type"
    test_string = "http-https-redirection"

    Result = CA.string_is_in_enumerated_property_list(profile_name,test_string)

    assert Result == True 

def test_string_is_in_enumerated_properties_list_returns_false_for_strings_not_in_list():

    profile_name = "anqp_nwk_auth_prof.anqp_nwk_auth_type.anqp_nwk_auth_type"
    test_string = "aruba-central"

    Result = CA.string_is_in_enumerated_property_list(profile_name,test_string)

    assert Result == False

def test_get_attribute_min_len_returns_correct_minimum_number_for_non_nested_string():

    profile_name = "ap_mesh_radio_prof.profile-name"
    
    expected = 1

    generated = CA.get_attribute_min_len(profile_name)

    assert expected == generated

def test_get_attribute_max_len_returns_correct_maximum_number_for_non_nested_string():

    profile_name = "ap_mesh_radio_prof.profile-name"
    
    expected = 256

    generated = CA.get_attribute_max_len(profile_name)

    assert expected == generated

def test_get_attribute_min_len_returns_correct_minimum_number_for_nested_string():

    profile_name = "ap_mesh_radio_prof.mesh_allowed_vlans.vlan-list"
    
    expected = 1

    generated = CA.get_attribute_min_len(profile_name)

    assert expected == generated

def test_get_attribute_max_len_returns_correct_maximum_number_for_nested_string():

    profile_name = "ap_mesh_radio_prof.mesh_allowed_vlans.vlan-list"
    
    expected = 256

    generated = CA.get_attribute_max_len(profile_name)

    assert expected == generated

def test_is_valid_string_returns_true_for_string_length_check():

    profile_name = "ap_mesh_radio_prof.profile-name"
    string = "some_profile_name"

    Result = CA.is_valid_string_or_number(profile_name,string)

    assert Result == True

def test_is_valid_string_returns_false_for_null_string_length_check():

    profile_name = "ap_mesh_radio_prof.profile-name"
    string = ""

    Result = CA.is_valid_string_or_number(profile_name,string)

    assert Result == False

def test_is_valid_string_returns_false_for_too_long_string_length_check():

    profile_name = "ap_mesh_radio_prof.profile-name"
    string = "a"*257

    Result = CA.is_valid_string_or_number(profile_name,string)

    assert Result == False

def test_is_valid_string_returns_true_for_nested_string_length_check():

    profile_name = "ap_mesh_radio_prof.mesh_allowed_vlans.vlan-list"
    string = "WIRELESS"

    Result = CA.is_valid_string_or_number(profile_name,string)

    assert Result == True

def test_is_valid_string_returns_false_for_null_nested_string_length_check():

    profile_name = "ap_mesh_radio_prof.mesh_allowed_vlans.vlan-list"
    string = ""

    Result = CA.is_valid_string_or_number(profile_name,string)

    assert Result == False

def test_is_valid_string_returns_false_for_too_long_nested_string_length_check():

    profile_name = "ap_mesh_radio_prof.mesh_allowed_vlans.vlan-list"
    string = "a"*257

    Result = CA.is_valid_string_or_number(profile_name,string)

    assert Result == False

def test_get_required_properties_returns_expected_required_list_when_it_exists():

    full_attribute_name = "ids_signature_prof.ids_condition_frame_type"

    expected = [
                        "control",
                        "deauth",
                        "frame_type_w_ssid",
                        "auth",
                        "assoc",
                        "disassoc",
                        "mgmt",
                        "data"
                    ]
    
    generated = CA.get_required_properties(full_attribute_name)

    for item in expected:
        assert item in generated

def test_get_required_properties_returns_empty_required_list_when_it_doesnt_exist():

    full_attribute_name = "ap_mesh_radio_prof.mesh_a_tx_rates"
    
    generated = CA.get_required_properties(full_attribute_name)

    assert len(generated) == 0

def test_property_is_in_properties_returns_true_for_a_existing_property():

    profile_name = "ap_mesh_radio_prof"
    attribute_name = "mesh_a_tx_rates"

    Result = CA.property_is_in_properties(profile_name,attribute_name,'6')

    assert Result == True

def test_property_is_in_properties_returns_false_for_a_nonexisting_property():

    profile_name = "ap_mesh_radio_prof"
    attribute_name = "mesh_a_tx_rates"

    Result = CA.property_is_in_properties(profile_name,attribute_name,'127')

    assert Result == False

def test_is_valid_string_or_number_returns_true_for_number_in_bound():

    profile_name = "ap_g_radio_prof.max_distance.maximum-distance"

    Result = CA.is_valid_string_or_number(profile_name,100,type="integer")

    assert Result == True

def test_is_valid_string_or_number_returns_false_for_number_too_big():

    profile_name = "ap_g_radio_prof.max_distance.maximum-distance"

    Result = CA.is_valid_string_or_number(profile_name,1000000,type="integer")

    assert Result == False

def test_is_valid_string_or_number_returns_false_for_number_too_small():

    profile_name = "ap_g_radio_prof.max_distance.maximum-distance"

    Result = CA.is_valid_string_or_number(profile_name,-1,type="integer")

    assert Result == False

def test_add_integer_attributes_to_profiles_correctly_adds_attributes_to_profiles():

    full_attribute_name = 'ap_multizone_prof.controller.num_vaps'

    profiles = [{'controller': {}}]

    attributes = ['15']

    expected = 15
    CA.add_integer_attribute_to_profiles(full_attribute_name,attributes,profiles)

    assert expected == profiles[0]['controller']['num_vaps']

def test_add_string_attributes_to_profiles_correctly_adds_attributes_to_profiles():

    full_attribute_name = 'ap_multizone_prof.profile-name'

    profiles = [{}]
    attributes = ['test_ap_multizone_prof']

    CA.add_string_attribute_to_profiles(full_attribute_name,attributes,profiles)

    assert 'test_ap_multizone_prof' == profiles[0]['profile-name']

def test_add_string_attribute_to_profiles_correctly_adds_nested_strings():

    full_attribute_name = 'ap_multizone_prof.ap_multizone_prof_clone.source'

    profiles = [{'ap_multizone_prof_clone':{}}]
    attributes = ['test_ap_multizone_prof']

    CA.add_string_attribute_to_profiles(full_attribute_name,attributes,profiles)

    assert 'test_ap_multizone_prof' == profiles[0]['ap_multizone_prof_clone']['source']

def test_add_boolean_attribute_to_profiles_correctly_adds_attribute():

    full_attribute_name = 'airmatch_ap_unfreeze.all-aps'

    profiles = [{}]
    attributes = ['True']

    CA.add_boolean_attribute_to_profiles(full_attribute_name,attributes,profiles)

    assert True == profiles[0]['all-aps']

def test_add_boolean_attribute_to_profiles_correctly_adds_nested_boolean_attribute():

    full_attribute_name = 'ssid_prof.opmode.wpa3-sae-aes'

    profiles = [{'opmode':{}}]
    attributes = ['True']

    CA.add_boolean_attribute_to_profiles(full_attribute_name,attributes,profiles)

    assert True == profiles[0]['opmode']['wpa3-sae-aes']

def test_add_object_attribute_to_profiles_correctly_adds_an_empty_object_attribute():

    full_attribute_name = 'ssid_prof.enable_ssid'

    profiles = [{}]
    attributes = ['True']

    CA.add_object_attribute_to_profiles(full_attribute_name,attributes,profiles)

    assert {} == profiles[0]['enable_ssid']

def test_add_object_attribute_to_profiles_correctly_adds_an_enumerated_string_correctly():

    full_attribute_name = 'ssid_prof.wmm_eap_ac'

    profiles = [{}]
    attributes = ['Voice']

    CA.add_entries_to_object_identifiers()
    CA.build_tables_columns_dict(Document('../Extension Testing Tables.docx').tables)
    CA.remove_column_headers_from_columns_table()
    CA.add_object_attribute_to_profiles(full_attribute_name,attributes,profiles)

    assert 'best-effort' == profiles[0]['wmm_eap_ac']['wmm_ac']

def test_add_object_attribute_to_profiles_adds_string_to_object_attribute():

    full_attribute_name = "ap_group.dot11a_prof"

    profiles = [{}]
    attributes = ['https://www.google.com']

    CA.add_entries_to_object_identifiers()
    CA.build_tables_columns_dict(Document('../Extension Testing Tables.docx').tables)
    CA.remove_column_headers_from_columns_table()
    CA.add_object_attribute_to_profiles(full_attribute_name,attributes,profiles)

    assert 'HQ' == profiles[0]['dot11a_prof']['profile-name']

def test_get_attribute_properties_returns_list_of_properties():

    full_attribute_name = "ap_mesh_radio_prof.mesh_a_tx_rates"

    expected = ['6','9','12','18','24','36','48','54']
    generated = CA.get_attribute_properties(full_attribute_name)

    assert expected == generated

def test_get_attribute_properties_returns_empty_list_for_attributes_without_properties():

    full_attribute_name = "ap_mesh_radio_prof.mesh_mcast_opt"

    expected = []
    generated = CA.get_attribute_properties(full_attribute_name)

    assert expected == generated

def test_add_object_attributes_to_profiles_adds_boolean_objects_correctly():

    full_attribute_name = "ids_general_prof.frame_types_for_rssi"

    profiles = [{}]
    attributes = ['Beacon,Probe,Management,Control']

    CA.add_object_attribute_to_profiles(full_attribute_name,attributes,profiles)

    assert {'ba':True,'pr':True,'mgmt':True,'ctrl':True} == profiles[0]['frame_types_for_rssi']

def test_add_object_attributes_to_profiles_adds_integer_boolean_objects_correctly():

    full_attribute_name = "ssid_prof.a_basic_rates"

    profiles = [{}]
    attributes = ['6,12,18,24,36']

    CA.add_object_attribute_to_profiles(full_attribute_name,attributes,profiles)

    assert {'6':True,'12':True,'18':True,'24':True,'36':True}

def test_add_object_attributes_to_profiles_raises_value_error_when_boolean_not_defined():

    full_attribute_name = "ids_general_prof.frame_types_for_rssi"

    profiles = [{}]
    attributes = ['Beacon,Probes,Management,Control']
    with pytest.raises(ValueError):
        CA.add_object_attribute_to_profiles(full_attribute_name,attributes,profiles)

def test_get_array_property_type_returns_correct_type():

    full_attribute_name = 'ap_group.virtual_ap'
    property = 'profile-name'

    type = CA.get_array_property_type(full_attribute_name,property)

    assert 'string' == type

def test_get_array_property_minimum_length_returns_correct_minimum():

    full_attribute_name = 'ap_group.virtual_ap.profile-name'

    min = CA.get_array_attribute_min_length(full_attribute_name)

    assert min == 1
    

def test_get_array_property_maximum_length_returns_correct_minimum():

    full_attribute_name = 'ap_group.virtual_ap.profile-name'

    max = CA.get_array_attribute_max_length(full_attribute_name)

    assert max == 256

def test_is_enumerated_array_property_returns_true_for_enumerated_property():

    full_attribute_name = 'usb_acl_prof.usb_access_rule.usb_vendor_list'
    
    assert True == CA.is_enumerated_array_property(full_attribute_name)

def test_is_enumerated_array_property_returns_false_for_non_enumerated_property():

    full_attribute_name = 'ids_rap_wml_server_prof.wml_table.profile-name'
    
    assert False == CA.is_enumerated_array_property(full_attribute_name)

def test_add_array_attribute_correctly_adds_attributes():

    full_attribute_name = "wlan_qos_prof.bw_alloc"

    CA.add_entries_to_object_identifiers()
    CA.build_tables_columns_dict(Document('../Extension Testing Tables.docx').tables)
    CA.remove_column_headers_from_columns_table()

    profiles = [{}]

    CA.add_array_attribute_to_profiles(full_attribute_name,profiles)

    assert [{'virtual-ap':'ChoiceAccess'},{'share':20}] == profiles[0]['bw_alloc']

def test_profile_is_an_attribute_of_current_profile_returns_true_for_properties():

    current_profile = 'ap_group'
    other_profile = 'dot11a_snd_prof'

    Result = CA.profile_is_an_attribute_of_current_profile(current_profile,other_profile)

    assert Result == True

def test_profile_is_an_attribute_of_current_profile_returns_false_for_nonproperties():

    current_profile = 'ap_group'
    other_profile = 'dank_prof'

    Result = CA.profile_is_an_attribute_of_current_profile(current_profile,other_profile)

    assert Result == False

def test_add_dependencies_to_table_columns_dict_adds_dependent_profiles():

    CA.add_entries_to_object_identifiers()
    CA.build_tables_columns_dict(Document('../Radio Testing Tables.docx').tables)
    CA.remove_column_headers_from_columns_table()
    profiles = CA.get_profiles_to_be_configured()
    CA.build_profiles_dependencies(profiles)

    expected = {"ap_g_radio_prof.profile-name":["CGR","DEP","KFD","PWK","SIG","Test"],
                "ap_a_radio_prof.profile-name":["CGR","DEP","KFD","PWK","SIG","Test"],
                "reg_domain_prof.profile-name":["CGR","DEP","KFD","PWK","SIG","Test"],
                "ap_g_radio_prof.eirp_min.eirp-min":["3 dBm", "3 dBm", "3 dBm", "3 dBm", "3 dBm", "3 dBm"],
                "ap_g_radio_prof.eirp_max.eirp-max":["7 dBm", "7 dBm", "7 dBm", "7 dBm", "7 dBm", "7 dBm"],
                "ap_a_radio_prof.eirp_min.eirp-min":["6 dBm", "6 dBm", "6 dBm", "6 dBm", "6 dBm", "6 dBm"],
                "ap_a_radio_prof.eirp_max.eirp-max":["10 dBm", "10 dBm", "10 dBm", "10 dBm", "10 dBm", "10 dBm"],
                "reg_domain_prof.valid_11b_channel.valid-11g-channel":["1, 6, 11", "1,6,11", "1, 6, 11", "1, 6, 11", "1, 6, 11", "1, 6, 11"],
                "reg_domain_prof.valid_11a_channel.valid-11a-channel":["36, 40, 44, 48, 52, 56, 60, 64, 149, 153, 157, 161,165", "36, 40, 44, 48, 52, 56, 60, 64, 149, 153, 157, 161,165", "36, 40, 44, 48, 52, 56, 60, 64, 149, 153, 157, 161,165", "36, 40, 44, 48, 52, 56, 60, 64, 149, 153, 157, 161,165", "36, 40, 44, 48, 52, 56, 60, 64, 149, 153, 157, 161,165", "36, 40, 44, 48, 149, 153, 157, 161,165" ],
                "ssid_prof.profile-name":["ChoiceAccess", "CorpAccess", "CorpTest", "GuestAccess", "MobiAccess"],
                "ssid_prof.essid.essid":["ChoiceAccess", "CorpAccess", "CorpTest", "GuestAccess", "MobiAccess"],
                "ssid_prof.g_basic_rates":["12", "12", "12", "12", "12"],
                "ssid_prof.g_tx_rates":["12, 18, 24, 36, 48, 54", "12, 18, 24, 36, 48, 54", "12, 18, 24, 36, 48, 54", "12, 18, 24, 36, 48, 54", "12, 18, 24, 36, 48, 54"],
                "ssid_prof.a_basic_rates":["12", "12", "12, 24", "12", "12"],
                "ssid_prof.a_tx_rates":["12, 18, 24, 36, 48, 54", "12, 18, 24, 36, 48, 54", "12, 18, 24, 36, 48, 54", "12, 18, 24, 36, 48, 54", "12, 18, 24, 36, 48, 54"],
                "ap_group.profile-name":["CGR","DEP","KFD","PWK","SIG","Test"],
                "ap_group.dot11g_prof.profile-name":["CGR_ap_g_radio_prof","DEP_ap_g_radio_prof","KFD_ap_g_radio_prof","PWK_ap_g_radio_prof","SIG_ap_g_radio_prof","Test_ap_g_radio_prof"],
                "ap_group.dot11a_prof.profile-name":["CGR_ap_a_radio_prof","DEP_ap_a_radio_prof","KFD_ap_a_radio_prof","PWK_ap_a_radio_prof","SIG_ap_a_radio_prof","Test_ap_a_radio_prof"],
                "ap_group.reg_domain_prof.profile-name":["CGR_reg_domain_prof","DEP_reg_domain_prof","KFD_reg_domain_prof","PWK_reg_domain_prof","SIG_reg_domain_prof","Test_reg_domain_prof"]}

    assert expected == CA.TABLE_COLUMNS

def test_add_dependencies_to_table_columns_dict_adds_new_profiles_to_profiles_to_be_configured():

    CA.add_entries_to_object_identifiers()
    CA.build_tables_columns_dict(Document('../Radio Testing Tables.docx').tables)
    CA.remove_column_headers_from_columns_table()
    profiles = CA.get_profiles_to_be_configured()
    CA.build_profiles_dependencies(profiles)

    assert 'dot11g_prof.profile-name' in profiles['ap_group'] and 'dot11a_prof.profile-name' in profiles['ap_group'] and 'reg_domain_prof.profile-name' in profiles['ap_group']                

def test_build_ordered_configuration_list_builds_list_correctly():

    CA.add_entries_to_object_identifiers()
    CA.build_tables_columns_dict(Document('../Radio Testing Tables.docx').tables)
    CA.remove_column_headers_from_columns_table()
    profiles = CA.get_profiles_to_be_configured()
    CA.build_profiles_dependencies(profiles)
    
    expected = [["ap_g_radio_prof.profile-name",
                "ap_g_radio_prof.eirp_min.eirp-min",
                "ap_g_radio_prof.eirp_max.eirp-max"],
                ["ap_a_radio_prof.profile-name",
                "ap_a_radio_prof.eirp_min.eirp-min",
                "ap_a_radio_prof.eirp_max.eirp-max"],
                ["reg_domain_prof.profile-name",
                "reg_domain_prof.valid_11b_channel.valid-11g-channel",
                "reg_domain_prof.valid_11a_channel.valid-11a-channel"],
                ["ap_group.profile-name",
                "ap_group.dot11g_prof.profile-name",
                "ap_group.dot11a_prof.profile-name",
                "ap_group.reg_domain_prof.profile-name"],
                ["ssid_prof.profile-name",
                 "ssid_prof.essid.essid",
                "ssid_prof.g_basic_rates",
                "ssid_prof.g_tx_rates",
                "ssid_prof.a_basic_rates",
                "ssid_prof.a_tx_rates"]]

    generated = CA.build_ordered_configuration_list(profiles)

    assert expected == generated

def test_get_api_endpoints_returns_correct_list_of_endpoints():

    ordered_profiles = [['ap_g_radio_prof.profile-name'],['ap_a_radio_prof.profile-name'],['reg_domain_prof.profile-name'],['ap_group.profile-name'],['ssid_prof.profile-name'],['virtual_ap_prof.profile-name']]
    expected = ['ap_g_radio_prof','ap_a_radio_prof','reg_domain_prof','ap_group','ssid_prof','virtual_ap_prof']

    generated = CA.get_list_of_api_endpoints(ordered_profiles)

    assert expected == generated

def test_add_src_dst_values_to_ace_object_adds_host_correctly():

    ace_values = ['host','1.1.1.1','192.168.1.1']
    ace_object = {}

    expected = {'src':'shost','sipaddr':'1.1.1.1'}

    CA.add_src_dst_values_to_ace_object(ace_values,ace_object,'s')

    assert expected == ace_object

def test_add_src_dst_values_to_ace_object_adds_network_correctly():

    ace_values = ['network', '192.168.1.1/29','other_value']
    ace_object = {}

    expected = {'src':'snetwork','snetwork':'192.168.1.1','snetmask':'255.255.255.248'}

    CA.add_src_dst_values_to_ace_object(ace_values,ace_object,'s')

    assert expected == ace_object

def test_add_src_dst_values_to_ace_object_adds_any_source_correctly():

    ace_values = ['any','other_value']
    ace_object = {}

    expected = {'src':'sany','sany':True}

    CA.add_src_dst_values_to_ace_object(ace_values,ace_object,'s')

    assert expected == ace_object

def test_add_src_dst_values_to_ace_object_adds_role_source_correctly():

    ace_values = ['role','CP-logon-role','other_value']
    ace_object = {}

    expected = {'src':'suserrole','surname':'CP-logon-role'}

    CA.add_src_dst_values_to_ace_object(ace_values,ace_object,'s')

    assert expected == ace_object

def test_add_src_dst_values_to_ace_object_adds_alias_source_correctly():

    ace_values = ['alias','CP-net-dest','other_value']
    ace_object = {}

    expected = {'src':'salias','srcalias':'CP-net-dest'}

    CA.add_src_dst_values_to_ace_object(ace_values,ace_object,'s')

    assert expected == ace_object

def test_add_src_dst_values_to_ace_object_adds_user_source_correctly():

    ace_values = ['user','CP-net-dest','other_value']
    ace_object = {}

    expected = {'src':'suser','suser':True}

    CA.add_src_dst_values_to_ace_object(ace_values,ace_object,'s')

    assert expected == ace_object

def test_add_action_values_to_acl_adds_permit_action_correctly():

    ace_values = ['permit', 'other_value']
    ace_object = {}

    expected = {'permit':True, 'action':'permit'}

    CA.add_action_values_to_ace_object(ace_values,ace_object)

    assert expected == ace_object

def test_add_action_values_to_acl_adds_deny_action_correctly():

    ace_values = ['deny', 'other_value']
    ace_object = {}

    expected = {'deny':True, 'action':'deny'}

    CA.add_action_values_to_ace_object(ace_values,ace_object)

    assert expected == ace_object

def test_add_action_values_to_acl_adds_dst_nat_ip():

    ace_values = ['dst-nat','ip','1.1.1.1','other_value']
    ace_object = {}

    expected = {'action':'dst-nat','dnatip':'1.1.1.1'}

    CA.add_action_values_to_ace_object(ace_values,ace_object)

    assert expected == ace_object

def test_add_action_values_to_acl_adds_dst_nat_name():

    ace_values = ['dst-nat','name','some_name.domain','other_value']
    ace_object = {}

    expected = {'action':'dst-nat','dnathostname':'some_name.domain'}

    CA.add_action_values_to_ace_object(ace_values,ace_object)

    assert expected == ace_object

def test_add_action_values_to_acl_adds_src_nat_without_pool_correctly():

    ace_values = ['src-nat','other_value']
    ace_object = {}

    expected = {'action':'src-nat','src-nat':True}

    CA.add_action_values_to_ace_object(ace_values,ace_object)

    assert expected == ace_object

def test_add_action_values_to_acl_adds_src_nat_with_pool_correctly():

    ace_values = ['src-nat','pool','srcnatpool','other_value']
    ace_object = {}

    expected = {'action':'src-nat','src-nat':True,'poolname':'srcnatpool'}

    CA.add_action_values_to_ace_object(ace_values,ace_object)

    assert expected == ace_object

def test_add_action_values_to_acl_adds_redirect_with_tunnel_id_correctly():

    ace_values = ['redirect','tunnel','10','other_value']
    ace_object = {}

    expected = {'action':'redir_opt','redir_opt':'tunnel','tunid':10}

    CA.add_action_values_to_ace_object(ace_values,ace_object)

    assert expected == ace_object

def test_add_action_values_to_acl_adds_redirect_with_tunnel_group_correctly():

    ace_values = ['redirect','tunnel-group','tungroup','other_value']
    ace_object = {}

    expected = {'action':'redir_opt','redir_opt':'tunnel-group','tungrpname':'tungroup'}

    CA.add_action_values_to_ace_object(ace_values,ace_object)

    assert expected == ace_object

def test_add_action_values_to_acl_adds_redirect_with_esi_group_with_direction_correctly():

    ace_values = ['redirect','group','esigroup','direction','forward','other_value']
    ace_object = {}

    expected = {'action':'redir_opt','redir_opt':'esi-group','group':'esigroup','dir':'forward'}

    CA.add_action_values_to_ace_object(ace_values,ace_object)

    assert expected == ace_object


def test_add_action_values_to_acl_adds_route_src_nat_correctly():

    ace_values = ['route','src-nat','esigroup','direction','forward','other_value']
    ace_object = {}

    expected = {'action':'route','src-nat-route':True}

    CA.add_action_values_to_ace_object(ace_values,ace_object)

    assert expected == ace_object

def test_add_action_values_to_acl_adds_route_dst_nat_with_ip_correctly():

    ace_values = ['route','dst-nat','ip','4.4.4.4','forward','other_value']
    ace_object = {}

    expected = {'action':'route','dst-nat-route':True,'routednatip':'4.4.4.4'}

    CA.add_action_values_to_ace_object(ace_values,ace_object)
    
    assert expected == ace_object


def test_add_action_values_to_acl_adds_route_dst_nat_with_name_correctly():

    ace_values = ['route','dst-nat','name','google.com','forward','other_value']
    ace_object = {}

    expected = {'action':'route','dst-nat-route':True,'routednathostname':'google.com'}

    CA.add_action_values_to_ace_object(ace_values,ace_object)
    
    assert expected == ace_object

def test_add_action_values_to_acl_returns_empty_array_without_extended_options():

    ace_values = ['redirect','group','esigroup','direction','forward']
    ace_object = {}

    generated = CA.add_action_values_to_ace_object(ace_values,ace_object)

    assert [] == generated

def test_add_extended_action_values_to_acl_adds_a_time_range_correctly():

    ace_values = ['time-range','after-lunch']
    ace_object = {}

    expected = {'trname':'after-lunch'}
    CA.add_extended_action_values_to_ace_object(ace_values,ace_object)

    assert expected == ace_object

def test_add_extended_action_values_to_acl_adds_a_queue_correctly():

    ace_values = ['queue','high']
    ace_object = {}

    expected = {'queue-type':'high', 'queue':True}
    CA.add_extended_action_values_to_ace_object(ace_values,ace_object)

    assert expected == ace_object

def test_add_extended_action_values_to_acl_adds_a_tos_value_correctly():

    ace_values = ['tos','46']
    ace_object = {}

    expected = {'tosstr':46}
    CA.add_extended_action_values_to_ace_object(ace_values,ace_object)

    assert expected == ace_object

def test_add_extended_action_values_to_acl_adds_a_dot1p_value_correctly():

    ace_values = ['priority-802.1p','7']
    ace_object = {}

    expected = {'prio8021p':7}
    CA.add_extended_action_values_to_ace_object(ace_values,ace_object)

    assert expected == ace_object


def test_add_extended_action_values_to_acl_adds_multiple_values_correctly():

    ace_values = ['priority-802.1p','7','time-range','lunch','log']
    ace_object = {}

    expected = {'prio8021p':7, 'trname':'lunch','log':True}
    CA.add_extended_action_values_to_ace_object(ace_values,ace_object)
    
    assert expected == ace_object

def test_add_service_values_to_acl_adds_any_service_correctly():

    ace_values = ['any','other_values']
    ace_object = {}

    expected = {'service_app':'service','service-any':True,'svc':'service-any'}
    CA.add_service_to_ace_object(ace_values,ace_object)

    assert expected == ace_object

def test_add_service_values_to_acl_adds_udp_with_port_correctly():

    ace_values = ['udp','44','other_values']
    ace_object = {}

    expected = {'service_app':'service','proto':'udp','svc':'tcp_udp','port1':44,'port':'range'}
    CA.add_service_to_ace_object(ace_values,ace_object)

    assert expected == ace_object

def test_add_service_values_to_acl_adds_udp_with_port_range_correctly():

    ace_values = ['udp','44','60','other_values']
    ace_object = {}

    expected = {'service_app':'service','proto':'udp','svc':'tcp_udp','port1':44,'port2':60,'port':'range'}
    CA.add_service_to_ace_object(ace_values,ace_object)

    assert expected == ace_object

def test_add_service_values_to_acl_adds_tcp_with_port_correctly():

    ace_values = ['tcp','45','other_values']
    ace_object = {}

    expected = {'service_app':'service','proto':'tcp','svc':'tcp_udp','port1':45,'port':'range'}
    CA.add_service_to_ace_object(ace_values,ace_object)

    assert expected == ace_object

def test_add_service_values_to_acl_adds_tcp_with_port_range_correctly():

    ace_values = ['tcp','45','70','other_values']
    ace_object = {}

    expected = {'service_app':'service','proto':'tcp','svc':'tcp_udp','port1':45,'port2':70,'port':'range'}
    CA.add_service_to_ace_object(ace_values,ace_object)

    assert expected == ace_object

def test_add_service_values_to_acl_adds_service_name_correctly():

    ace_values = ['svc-some-service','other_values']
    ace_object = {}

    expected = {'service_app':'service','svc':'service-name','service-name':'svc-some-service'}
    CA.add_service_to_ace_object(ace_values,ace_object)

    assert expected == ace_object

def test_add_service_values_to_acl_adds_protocol_correctly():

    ace_values = ['17','other_values']
    ace_object = {}

    expected = {'service_app':'service','svc':'protocol','protocol':17}
    CA.add_service_to_ace_object(ace_values,ace_object)

    assert expected == ace_object

def test_add_aces_to_acl_produces_acl_correctly():

    ace = "user host 10.10.10.100 any permit log"
    
    expected = {'suser':True,'src':'suser','dipaddr':'10.10.10.100','dst':'dhost','service_app':'service','svc':'service-any','service-any':True,'action':'permit','permit':True,'log':True}
    generated = CA.translate_ace_to_api_values(ace)

    assert expected == generated

def test_add_aces_to_acl_produces_acl_correctly_02():

    ace = "network 192.168.1.0/24 any udp 68 deny log time-range lunch mirror"
    
    expected = {'snetwork':'192.168.1.0','src':'snetwork','snetmask':'255.255.255.0','dany':True,'dst':'dany','service_app':'service','svc':'tcp_udp','proto':'udp','port1':68,'port':'range','action':'deny','deny':True,'log':True,'trname':'lunch','mirror':True}
    generated = CA.translate_ace_to_api_values(ace)

    assert expected == generated

def test_get_column_errors_processes_single_integer_value_in_range_correctly():

    doc = Document()
    table = doc.add_table(cols=1,rows=2)
    cells = table.columns[0].cells
    cells[0].text = 'Int VLAN ID'
    cells[1].text = '20'

    CA.add_entries_to_object_identifiers()
    CA.build_tables_columns_dict([table])
    generated = CA.get_column_errors()

    assert [] == generated

def test_get_column_errors_processes_single_integer_value_out_of_range_correctly():

    doc = Document()
    table = doc.add_table(cols=2,rows=2)
    cells = table.columns[0].cells
    cells[0].text = 'Node'
    cells[1].text = '/md/America/East'
    cells = table.columns[1].cells
    cells[0].text = 'Int VLAN ID'
    cells[1].text = '4098'

    CA.add_entries_to_object_identifiers()
    CA.build_tables_columns_dict([table])
    generated = CA.get_column_errors()

    assert [{"Int VLAN ID":[[0, 0,'4098']]}] == generated

def test_get_column_errors_processes_single_string_value_correctly():

    doc = Document()
    table = doc.add_table(cols=2,rows=2)
    cells = table.columns[0].cells
    cells[0].text = 'Node'
    cells[1].text = '/md/America/East'
    cells = table.columns[1].cells
    cells[0].text = 'DHCP Pool Name'
    cells[1].text = 'Some_Name'

    CA.add_entries_to_object_identifiers()
    CA.build_tables_columns_dict([table])
    generated = CA.get_column_errors()

    assert [] == generated

def test_get_column_errors_processes_multiple_string_values_correctly():

    doc = Document()
    table = doc.add_table(cols=2,rows=2)
    cells = table.columns[0].cells
    cells[0].text = 'Node'
    cells[1].text = '/md/America/East'
    cells = table.columns[1].cells
    cells[0].text = 'SG Server Name'
    cells[1].text = 'Server1, Server2, Server3'

    CA.add_entries_to_object_identifiers()
    CA.build_tables_columns_dict([table])
    generated = CA.get_column_errors()

    assert [] == generated

def test_get_column_errors_catches_invalid_string_in_multiple_string_values_correctly():

    doc = Document()
    table = doc.add_table(cols=2,rows=2)
    cells = table.columns[0].cells
    cells[0].text = 'Node'
    cells[1].text = '/md/America/East'
    cells = table.columns[1].cells
    cells[0].text = 'SG Server Name'
    cells[1].text = 'Server1'*40 + ', Server2, Server3'

    CA.add_entries_to_object_identifiers()
    CA.build_tables_columns_dict([table])
    generated = CA.get_column_errors()

    assert [{'SG Server Name':[[0, 0, 'Server1'*40]]}] == generated

def test_get_column_errors_catches_invalid_strings_in_multiple_string_values_correctly():

    doc = Document()
    table = doc.add_table(cols=2,rows=2)
    cells = table.columns[0].cells
    cells[0].text = 'Node'
    cells[1].text = '/md/America/East'
    cells = table.columns[1].cells
    cells[0].text = 'SG Server Name'
    cells[1].text = 'Server1'*40 + ', Server2'
    cells[1].text += ', '+'Server3'*40

    CA.add_entries_to_object_identifiers()
    CA.build_tables_columns_dict([table])
    generated = CA.get_column_errors()

    assert [{'SG Server Name':[[0, 0, 'Server1'*40],[0, 2, 'Server3'*40]]}] == generated

def test_string_is_in_enumerated_list_checks_for_attributes_two_deep():

    full_attribute_name = 'radius_attr.attr_type'
    
    Result = CA.string_is_in_enumerated_property_list(full_attribute_name,'ipaddr')

    assert Result == True

def test_string_is_in_enumerated_list_checks_for_attributes_three_deep():

    full_attribute_name = 'mgmt_auth_profile.mgmt_default_role.aaa_auth_mgmt_default_role'
    
    Result = CA.string_is_in_enumerated_property_list(full_attribute_name,'guest-provisioning')

    assert Result == True

def test_get_column_skips_empty_object_attributes():

    doc = Document()
    table = doc.add_table(cols=2,rows=2)
    cells = table.columns[0].cells
    cells[0].text = 'Node'
    cells[1].text = '/md/America/East'
    cells = table.columns[1].cells
    cells[0].text = 'MFP/PMF'
    cells[1].text = 'True'

    CA.add_entries_to_object_identifiers()
    CA.build_tables_columns_dict([table])
    generated = CA.get_column_errors()

    assert [] == generated

def test_get_column_errors_skips_boolean_attributes():

    doc = Document()
    table = doc.add_table(cols=2,rows=2)
    cells = table.columns[0].cells
    cells[0].text = 'Node'
    cells[1].text = '/md/America/East'
    cells = table.columns[1].cells
    cells[0].text = 'MFP/PMF'
    cells[1].text = 'True'

    CA.add_entries_to_object_identifiers()
    CA.build_tables_columns_dict([table])
    generated = CA.get_column_errors()

    assert [] == generated

def test_is_valid_object_returns_true_booleans_in_object():

    full_attribute_name = 'ssid_prof.a_basic_rates'
    attribute = '6'

    Result = CA.is_valid_object(attribute,full_attribute_name)

    assert Result == True

def test_is_valid_object_returns_true_for_empty_objects():

    full_attribute_name = 'ssid_prof.okc_enable'
    attribute = 'True'

    Result = CA.is_valid_object(attribute,full_attribute_name)

    assert Result == True

def test_add_5ghz_width_adds_40mhz_channels_correctly():

    doc = Document()
    table = doc.add_table(cols=3,rows=2)
    cells = table.columns[0].cells
    cells[0].text = 'Node'
    cells[1].text = '/md/America/East'
    cells = table.columns[1].cells
    cells[0].text = '5 GHz Channels'
    cells[1].text = '36, 40, 44, 48, 52, 56, 60, 64, 149, 153, 157, 161, 165'
    cells = table.columns[2].cells
    cells[0].text = '5 GHz Channel Width'
    cells[1].text = '40 MHz'

    expected = {
                'profile-name':'some-name',
                'valid_11a_40mhz_chan_pair_nd': {'valid-11a-40mhz-channel-pair-nd':['36-40','44-48','52-56','60-64','149-153','157-161']}
    }

    CA.add_entries_to_object_identifiers()
    CA.build_tables_columns_dict(doc.tables)
    CA.remove_column_headers_from_columns_table()
    generated = CA.add_attributes_to_profiles('reg_domain_prof.channel_width.width',[],[{'profile-name':'some-name'}])

    assert expected == generated[0]

def test_add_5ghz_width_adds_80mhz_channels_correctly():

    doc = Document()
    table = doc.add_table(cols=3,rows=2)
    cells = table.columns[0].cells
    cells[0].text = 'Node'
    cells[1].text = '/md/America/East'
    cells = table.columns[1].cells
    cells[0].text = '5 GHz Channels'
    cells[1].text = '36, 40, 44, 48, 52, 56, 60, 64, 149, 153, 157, 161, 165'
    cells = table.columns[2].cells
    cells[0].text = '5 GHz Channel Width'
    cells[1].text = '80 MHz'

    expected = {
                'profile-name':'some-name',
                'valid_11a_80mhz_chan_group': {'valid-11a-80mhz-channel-group':['36-48','52-64','149-161']}
    }

    CA.add_entries_to_object_identifiers()
    CA.build_tables_columns_dict(doc.tables)
    CA.remove_column_headers_from_columns_table()
    generated = CA.add_attributes_to_profiles('reg_domain_prof.channel_width.width',[],[{'profile-name':'some-name'}])

    assert expected == generated[0]

def test_add_5ghz_width_adds_160mhz_channels_correctly():

    doc = Document()
    table = doc.add_table(cols=3,rows=2)
    cells = table.columns[0].cells
    cells[0].text = 'Node'
    cells[1].text = '/md/America/East'
    cells = table.columns[1].cells
    cells[0].text = '5 GHz Channels'
    cells[1].text = '36, 40, 44, 48, 52, 56, 60, 64, 100, 104, 108, 112, 116, 120, 124, 128, 132, 136, 140, 144, 149, 153, 157, 161, 165'
    cells = table.columns[2].cells
    cells[0].text = '5 GHz Channel Width'
    cells[1].text = '160 MHz'

    expected = {
                'profile-name':'some-name',
                'valid_11a_160mhz_chan_group': {'valid-11a-160mhz-channel-group':['36-64','100-128','132-161']}
    }

    CA.add_entries_to_object_identifiers()
    CA.build_tables_columns_dict(doc.tables)
    CA.remove_column_headers_from_columns_table()
    generated = CA.add_attributes_to_profiles('reg_domain_prof.channel_width.width',[],[{'profile-name':'some-name'}])

    assert expected == generated[0]

def test_add_addresses_to_dhcp_pool_adds_dns_addresses_correctly():

    doc = Document()
    table = doc.add_table(cols=1,rows=2)
    cells = table.columns[0].cells
    cells[0].text = 'DHCP Pool DNS IP'
    cells[1].text = '192.168.1.21, 192.168.2.22, 192.168.3.33'

    expected = {'pool-name':'some-name','ip_dhcp_pool_cfg__dns':{'address1':'192.168.1.21','address2':'192.168.2.22','address3':'192.168.3.33'}}
    
    CA.add_entries_to_object_identifiers()
    CA.build_tables_columns_dict(doc.tables)
    CA.remove_column_headers_from_columns_table()
    generated = CA.add_attributes_to_profiles('ip_dhcp_pool_cfg.ip_dhcp_pool_cfg__dns.address1',[],[{'pool-name':'some-name'}])

    assert expected == generated[0]

def test_add_addresses_to_dhcp_pool_adds_dft_rtr_addresses_correctly():

    doc = Document()
    table = doc.add_table(cols=1,rows=2)
    cells = table.columns[0].cells
    cells[0].text = 'DHCP Pool Default Router'
    cells[1].text = '192.168.1.21, 192.168.2.22, 192.168.3.33'

    expected = {'pool-name':'some-name','ip_dhcp_pool_cfg__def_rtr':{'address':'192.168.1.21','address2':'192.168.2.22','address3':'192.168.3.33'}}
    
    CA.add_entries_to_object_identifiers()
    CA.build_tables_columns_dict(doc.tables)
    CA.remove_column_headers_from_columns_table()
    generated = CA.add_attributes_to_profiles('ip_dhcp_pool_cfg.ip_dhcp_pool_cfg__def_rtr.address',[],[{'pool-name':'some-name'}])

    assert expected == generated[0]

def test_add_node_info_to_object_identifier_uses_empty_node_names_correctly():

    doc = Document()
    table = doc.add_table(cols=2,rows=5)
    cells = table.columns[0].cells
    cells[0].text = 'Node'
    cells[1].text = '/md/America/West'
    cells[2].text = ''
    cells[3].text = ''
    cells[4].text = '/md/America/East'
    cells = table.columns[1].cells
    cells[0].text = 'ESSID'
    cells[1].text = 'WIRELESS'
    cells[2].text = 'EK-CORP'
    cells[3].text = 'EK-BYOD'
    cells[4].text = 'Guest'

    expected = {'ssid_prof.profile-name':['ESSID','WIRELESS,/md/America/West','EK-CORP,/md/America/West','EK-BYOD,/md/America/West','Guest,/md/America/East'],
                'ssid_prof.essid.essid':['ESSID','WIRELESS,/md/America/West','EK-CORP,/md/America/West','EK-BYOD,/md/America/West','Guest,/md/America/East']}

    CA.add_entries_to_object_identifiers()
    CA.build_tables_columns_dict(doc.tables)
    
    assert CA.TABLE_COLUMNS == expected

def test_add_vlan_name_id_association_adds_name_association_correctly():

    doc = Document()
    table = doc.add_table(cols=3,rows=2)
    cells = table.columns[0].cells
    cells[0].text = 'Node'
    cells[1].text = '/md/America/West'
    cells = table.columns[1].cells
    cells[0].text = 'VLAN Name'
    cells[1].text = 'BYOD'
    cells = table.columns[2].cells
    cells[0].text = 'VLAN ID'
    cells[1].text = '20'

    CA.add_entries_to_object_identifiers()
    CA.build_tables_columns_dict(doc.tables)
    CA.remove_column_headers_from_columns_table()
    generated = CA.add_attributes_to_profiles('vlan_name.name',[],[{}])

    expected = [{'name':'BYOD,/md/America/West'},{'name':'BYOD,/md/America/West', 'vlan-ids':'20'}]

    assert expected == generated 