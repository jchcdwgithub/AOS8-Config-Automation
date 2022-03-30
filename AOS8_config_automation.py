import re
import math
import api
import data_structures
from docx import Document
from pprint import pprint

API_REF = api.get_API_JSON_files()
TABLE_COLUMNS = {}
OBJECT_IDENTIFIERS = {}
DEVICE_DICTIONARY = {}

def build_hierarchy(full_path):
  """ Given a path i.e /md/HQ/DC or HQ/DC, create all the nodes in the path after /md. """
  
  node_path = full_path.split('/')
  if '' == node_path[0]: node_path = node_path[1:]
  if 'md' == node_path[0].lower(): node_path = node_path[1:]
  node_paths = []

  partial_path = '/md'
  for path in node_path:
    partial_path += '/' + path
    node_paths.append(partial_path)

  return node_paths

def get_list_of_api_endpoints(ordered_profiles):
  """ Given an ordered list of profiles to configure, extract the endpoint information and return them as a list. """

  api_endpoints = []

  for profile_group in ordered_profiles:
    profile_parts = profile_group[0].split('.')
    api_endpoint = profile_parts[0]
    if api_endpoint not in api_endpoints:
      api_endpoints.append(api_endpoint)
  
  return api_endpoints

def push_profiles_to_network(api_endpoints, profiles):
  """ Given a list of profiles, push them to the network. """

  for api_endpoint,profile_list in zip(api_endpoints,profiles):
    current_node = ''
    for profile in profile_list:
      node = profile.pop('node')
      if node != current_node:
        if current_node == '':
          current_node = node
        else:
          proceed = input('Must commit all chnages to current node to push configuration to other nodes. Commit changes to current node? (y/n)')
          if proceed == 'y':
            print(f"Saving configuration...")
            response = api.write_mem(path=current_node)
            if response.status_code != 200:
              print(f"Encountered problem when trying to write memory: HTTP Status Code: {response.status_code}. Aborting...")
              exit()
            else:
               current_node = node      
          else:
            print(f"Aborting...")
            exit()
      print(f"Pushing configuration object {api_endpoint} to {node}. Profile details:")
      pprint(profile)
      proceed = input("Proceed? (y/n)")
      if proceed == 'y':
        response = api.call_api(api_endpoint,config_path=node,data=profile,post=True)
        if (response.status_code == 200 and response.json()['_global_result']['status'] != 0) or response.status_code != 200:
          if response.status_code != 200:
            print(f"Something went wrong while trying to communicate with the MM. HTTP Status Code: {response.status_code}. Exiting...")
            #exit()
          else:
            print(f"One or more of the attributes in the configuration profile {profile} failed to push:")
            pprint(response.json())
        else:
          print(f'Successfully pushed all configuration profiles of type {api_endpoint} to {node}.')
      else:
        print('Continuing in test environment...')
        #exit()
    print(f"Saving configurations before moving on to the next set of profiles...")
    response = api.write_mem(path=node)
    if response.status_code != 200:
      print(f"Encountered problem when trying to write memory. HTTP Status Code: {response.status_code}. Aborting...")
      #exit()
    else:
      print("Saved configurations.")
  print(f'All configuration pushed to the network.')

def build_profiles_from_ordered_list(ordered_profiles):
  """ Given an ordered list of profiles to configure, build the profiles from the TABLE_COLUMNS information. """

  profiles = []
  for profile in ordered_profiles:
    current_profiles = []
    profile_type = profile[0].split('.')[0]

    for prof_attribute in profile:
      if prof_attribute == OBJECT_IDENTIFIERS[profile_type]:
        profile_names = TABLE_COLUMNS[prof_attribute]
        suffixed_names = []

        for profile_name in profile_names:
          if len(profile_name.split('%')) < 2:
            node = api.DEFAULT_PATH
            if profile_name != '':
              if '_prof' in profile_type or '_group' in profile_type:
                profile_name += f'_{profile_type}'
          else:
            name,node = profile_name.split('%')
            if name != '':
              if '_prof' in profile_type or '_group' in profile_type:
                profile_name = f'{name}_{profile_type}'
              else:
                profile_name = name
            else:
              profile_name = name
          suffixed_names.append(profile_name)
          current_profiles.append({'node':node})

        TABLE_COLUMNS[prof_attribute] = suffixed_names
    for profile_attribute in profile:
      add_attributes_to_profiles(profile_attribute,TABLE_COLUMNS[profile_attribute],current_profiles)
      if profile_attribute not in SPECIAL_COLUMNS:
        remove_empty_objects_that_are_not_booleans(profile_attribute,TABLE_COLUMNS[profile_attribute],current_profiles)
        current_profiles = remove_empty_string_objects(current_profiles)
      current_profiles = remove_node_only_profiles(current_profiles)

    profiles.append(current_profiles)
  
  add_vlan_name_association(profiles,ordered_profiles)

  return profiles

def remove_empty_string_objects(current_profiles):
  """ Removes any objects that might have empty strings. These will raise an incomplete configuration at runtime. """
  
  returned_profiles = []
  for profile in current_profiles:
    profile_copy = profile.copy()
    for attribute in profile:
      if isinstance(profile[attribute],list):
        for object in profile[attribute]:
          list_copy = []
          for key in object:
            if object[key] != '':
              list_copy.append({key:object[key]})
          if len(list_copy) != 0:
            profile_copy[attribute] = list_copy
          else:
            profile_copy.pop(attribute)
      else:
       if profile[attribute] == '':
         profile_copy.pop(attribute)
    returned_profiles.append(profile_copy)
  
  return returned_profiles

def remove_node_only_profiles(profiles):
  """ Removes profiles that only have a node attribute after having added attributes to the profile. """

  non_empty_profiles = []
  for profile in profiles:
    if len(profile.keys()) != 1:
      non_empty_profiles.append(profile)
  
  return non_empty_profiles

def make_ordered_config_list(all_profiles):

  added_to_ordered_list = set()
  ordered_config_list = []
  for profile in all_profiles:
    if profile not in added_to_ordered_list:
      rec_build_ordered_config_list(profile, all_profiles, ordered_config_list,added_to_ordered_list)

  return ordered_config_list

def rec_build_ordered_config_list(current_profile, all_profiles, order_config_list, added):

  if current_profile not in all_profiles:
    print(f"{current_profile} is a nested object but does not have any corresponding configuration data. If it is a system defined/default object then ignore this message. Otherwise, add configuration for this object or the configuration will fail to push completely.")
    return order_config_list
  for attribute in all_profiles[current_profile]:
    attr_name = attribute.split('.')[0]
    if attribute_is_api_object(attr_name):
      if attr_name in data_structures.NESTED_DICT and current_profile != 'vlan_id':
        attr_name = data_structures.NESTED_DICT[attr_name][0]
      rec_build_ordered_config_list(attr_name, all_profiles, order_config_list, added)

  if current_profile not in added:  
    added.add(current_profile)
    return add_profile_to_ordered_configuration_list(current_profile, all_profiles, order_config_list)
  else:
    return order_config_list
          
def build_ordered_configuration_list(profiles_to_configure):
  """ Find objects that are dependent on each other in the TABLE_COLUMNS dictionary and add object attributes to be configured. """

  ordered_configuration_list = []
  added_objects = set()

  for profile in profiles_to_configure:
    for attribute in profiles_to_configure[profile]:
      attr_name = attribute.split('.')[0]
      if attribute_is_api_object(attr_name):
        if attr_name not in profiles_to_configure:
          attr_name = get_dependency_object_name(attr_name)
        if attr_name not in added_objects:
          add_profile_to_ordered_configuration_list(profile,profiles_to_configure,ordered_configuration_list,inner_profile=attr_name,added=added_objects)
          added_objects.add(attr_name)
    if profile not in added_objects:
      add_profile_to_ordered_configuration_list(profile,profiles_to_configure,ordered_configuration_list)
      added_objects.add(profile)
  
  return ordered_configuration_list

def add_profile_to_ordered_configuration_list(profile,profiles_to_configure,ordered_configuration_list):
  """ Adds the profile and all its attributes to be configured to the ordered list. """

  grouped_attributes = []
  for attribute in profiles_to_configure[profile]:
    grouped_attributes.append(f'{profile}.{attribute}')
  ordered_configuration_list.append(grouped_attributes)
  
def is_nested_attribute(attribute):
  """ Returns True if the attribute is a proper object in the API. """
  
  for object_name in data_structures.DEPENDENCY_DICT:
    if data_structures.DEPENDENCY_DICT[object_name] == attribute:
      return True
  
  if is_nested_profile(attribute):
    return True
  
  return False

def get_dependency_object_name(attribute):
  """ Given an attribute of an object, return the object name of the attribute if it exists. """
  
  for object_name in data_structures.DEPENDENCY_DICT:
    if data_structures.DEPENDENCY_DICT[object_name] == attribute:
      return object_name
  
  return ''

def get_object_name(object):
  """ Given an API object, return the ID for the object. """
  
  for attribute in object:
    if attribute in OBJECT_IDENTIFIERS:
      return attribute

def add_attributes_to_profiles(full_attribute_name,attributes,profiles):
  """ Checks against the API that the attributes in the column are correct and adds them to the provided profiles. """
  
  if full_attribute_name in SPECIAL_COLUMNS:
    profiles = SPECIAL_COLUMNS[full_attribute_name](profiles)
    return profiles
  elif len(full_attribute_name.split('.')) < 3:
    attribute_type = get_attribute_type(full_attribute_name)
  else:
    prof_name,attribute,_ = full_attribute_name.split('.')
    attribute_type = get_attribute_type(prof_name+'.'+attribute)
    if attribute_type == 'object':
      if attribute not in profiles[0].keys():
        add_object_attribute_to_profiles(prof_name+'.'+attribute,[],profiles)
        return profiles
      else:
        attribute_type = get_attribute_type(full_attribute_name)
    elif attribute_type == 'array':
      add_array_attribute_to_profiles(prof_name+'.'+attribute,profiles)
      return profiles
    
  if attribute_type == 'integer':
    add_integer_attribute_to_profiles(full_attribute_name,attributes,profiles)
  elif attribute_type == 'string':
    add_string_attribute_to_profiles(full_attribute_name,attributes,profiles)
  elif attribute_type == 'object':
    add_object_attribute_to_profiles(full_attribute_name,attributes,profiles)
  else:
    add_boolean_attribute_to_profiles(full_attribute_name,attributes,profiles)

  return profiles

def add_integer_attribute_to_profiles(full_attribute_name,attributes,profiles):
  """ Add the integer attribute to the provided profiles. """
  if len(full_attribute_name.split('.')) < 3:
    full_attribute_name += '.'
  
  prof_name,attribute_name,property_name = full_attribute_name.split('.')

  for profile,attribute in zip(profiles,attributes):
    attribute_number = extract_numbers_from_attribute(attribute)[0] 
    if attribute_number.isdigit():
      attribute_value = int(attribute_number)
      if is_valid_string_or_number(full_attribute_name,attribute_value,type="integer"):
        if property_name != '':
          profile[attribute_name][property_name] = attribute_value
        else:
          profile[attribute_name] = attribute_value
      else:
        raise ValueError(f'Invalid value. Number not within acceptable range: {prof_name} {attribute_name}')
    elif attribute == '':
      continue
    else:
      raise ValueError(f'Attribute {attribute_name} in {prof_name} must be a number.')      

def add_string_attribute_to_profiles(full_attribute_name,attributes,profiles):
  """ Add the string attributes to the provided profiles. """

  if len(full_attribute_name.split('.')) < 3:
    full_attribute_name += '.'
  
  prof_name,attribute_name,property_name = full_attribute_name.split('.')

  if property_name != '' and is_enumerated_property(full_attribute_name):
      for profile,attribute in zip(profiles,attributes):
        try:
          api_string = data_structures.BOOLEAN_DICT[attribute]
        except KeyError:
          print(f'Value {attribute} not defined in BOOLEAN_DICT when building {full_attribute_name}. Please add and try again.')
          exit()
        if string_is_in_enumerated_property_list(full_attribute_name,api_string):
          profile[attribute_name][property_name] = api_string
        else:
          raise ValueError(f'Invalid value {attribute} in {prof_name}. Fix and try again.')
  else:
    for profile,attribute in zip(profiles,attributes):
      if attribute != '':
        if is_valid_string_or_number(full_attribute_name,attribute):
          if property_name != '':
           profile[attribute_name][property_name] = attribute
          else:
            profile[attribute_name] = attribute
        else:
          exit()

def add_wide_5ghz_channels(profiles):
  """ Processes the 5ghz channel width property and produces the list of bonded channels for the corresponding bonded channel attribute
      in the SSID profile."""

  width_dict = {'40mhz':False,'80mhz':False,'160mhz':False}
  widths_column = TABLE_COLUMNS['reg_domain_prof.channel_width.width']

  for width in widths_column:
    if width in data_structures.BOOLEAN_DICT:
      width_value = data_structures.BOOLEAN_DICT[width]
      width_dict[width_value] = True
  
  if width_dict['40mhz'] or width_dict['80mhz'] or width_dict['160mhz']:
    try:
      wide_20 = TABLE_COLUMNS['reg_domain_prof.valid_11a_channel.valid-11a-channel']
    except KeyError:
      print('Must define an allowed 20 MHz channel column to use the 5 GHz Channel Width column.')
      exit()
  
  if width_dict['40mhz']:
    wide_40 = build_xmhz_channel_width_profiles('40mhz',widths_column,wide_20)
    for width_profile,profile in zip(wide_40,profiles):
      if width_profile != '':
        profile['valid_11a_40mhz_chan_pair_nd'] = width_profile

  if width_dict['80mhz']:
    wide_80 = build_xmhz_channel_width_profiles('80mhz',widths_column,wide_20)
    for width_profile,profile in zip(wide_80,profiles):
      if width_profile != '':
        profile['valid_11a_80mhz_chan_group'] = width_profile
  
  if width_dict['160mhz']:
    wide_160 = build_xmhz_channel_width_profiles('160mhz',widths_column,wide_20)
    for width_profile,profile in zip(wide_160,profiles):
      if width_profile != '':
        profile['valid_11a_160mhz_chan_group'] = width_profile

  return profiles

def build_xmhz_channel_width_profiles(channel_width,widths_column,wide_20):
  """ Returns a set of profiles given the channel_width and widths_column. channel_width is 40mhz, 80mhz or 160mhz. """

  chan_width_names = {'40mhz':['valid-11a-40mhz-channel-pair-nd',2],
                      '80mhz':['valid-11a-80mhz-channel-group',4],
                      '160mhz':['valid-11a-160mhz-channel-group',8]}
  
  prop_name = chan_width_names[channel_width][0]
  step = chan_width_names[channel_width][1]
  wide_xmhz = []
  pairs = make_pairs(wide_20,step)
  for width,pair in zip(widths_column,pairs):
    if data_structures.BOOLEAN_DICT[width] == channel_width:
      wide_xmhz.append({prop_name:pair})
    else:
      wide_xmhz.append('')
  
  return wide_xmhz


def make_pairs(wide_20,step):
  """ Given an array of comma separated 20 MHz channels, return a list of x MHz pairs where x is defined by
      the step size. 2 is 40 MHz, 4 is 80 MHz, 8 is 160 MHz. """
  
  pairs = []

  for channels in wide_20:
    channels_20 = channels.replace(' ','').split(',')
    current = 0
    next = step-1
    ties = []
    while current < len(channels_20)-1 and next < len(channels_20):
      ties.append(f"{channels_20[current]}-{channels_20[next]}")
      current += step
      next += step
    pairs.append(ties)

  return pairs   

def add_acls_to_role_profiles(profiles,attributes=[]):
  """ Special method for adding ACLs to user roles. """
  
  for attribute,profile in zip(attributes,profiles):
    acls = attribute.replace(', ',',').split(',')
    for acl in acls:
      try:
        _,acl_name = acl.split(' ')
      except:
        print('ACLs must be specified as either std, ext or session ACLS. ex. session ACL1, session ACL2, etc.')
        exit()
      profile['role__acl'].append({'pname':acl_name})
  
  return profiles

def add_addresses_to_dhcp_pool(profiles):
  """ DNS and DHCP addresses have special names in the API when adding more than one address.
      address:DNS1, address2:DNS2, etc. Up to 6 addresses can be added to the arrays. The attributes
      list contains the name of the addresses to configure: dns, dft_rtr, etc. """
  
  attributes = ['dns','def_rtr']

  for attribute in attributes:
    attr_name = f'ip_dhcp_pool_cfg.ip_dhcp_pool_cfg__{attribute}.address'
    if attribute == 'dns':
      attr_name += '1'
    if attr_name in TABLE_COLUMNS:
      addresses = TABLE_COLUMNS[attr_name]
  
      for profile,address in zip(profiles,addresses):
        address_list = address.replace(', ',',').split(',')
        profile[attr_name.split('.')[1]] = {}
        for number,address in enumerate(address_list,start=1):
          if number == 1 and attribute == 'def_rtr':
            address_name = 'address'
          else:
            address_name = f'address{number}'
          profile[f'ip_dhcp_pool_cfg__{attribute}'][address_name] = address
  
  return profiles

def add_lease_to_dhcp_pool(profiles):
  """ DHCP leases require var1 to var3 to be pushed together or else the attribute will be an incomplete command and will not push. """

  pool_names = TABLE_COLUMNS['ip_dhcp_pool_cfg.pool_name']
  base_name = 'ip_dhcp_pool_cfg.ip_dhcp_pool_cfg__lease.var'
  var_names = [f"{base_name}{num}" for num in ['1','2','3']]
  for var_name in var_names:
    if var_name not in TABLE_COLUMNS:
      TABLE_COLUMNS[var_name] = [0 for _ in pool_names]
  for index,table_column in enumerate([TABLE_COLUMNS[var_names[0]],TABLE_COLUMNS[var_names[1]],TABLE_COLUMNS[var_names[2]]]):
    new_column = []
    for var in table_column:
      if var == '':
        new_column.append(0)
      else:
        new_column.append(int(var))
    TABLE_COLUMNS[var_names[index]] = new_column
  for var1,var2,var3,profile in zip(TABLE_COLUMNS[var_names[0]],TABLE_COLUMNS[var_names[1]],TABLE_COLUMNS[var_names[2]],profiles):
    profile['ip_dhcp_pool_cfg__lease'] = {'var1':var1, 'var2':var2, 'var3':var3}
  return profiles 

def add_dns_server_addresses(profiles):
  """ There can be up to three DNS servers added to a controller. Since the feature is also an identifier for the name server,
      the processing happens here. """

  dns_addresses = TABLE_COLUMNS['ip_name_server.address']
  
  for dns_address,profile in zip(dns_addresses,profiles):
    addresses = dns_address.replace(' ','').split(',')
    if len(addresses) > 3:
      print('Only three DNS servers can be defined per controller. Using first three IPs provided...')
      addresses = addresses[:3]
    original = True
    new_profile = profile.copy()
    for address in addresses:
      if is_valid_ip_address(address):
        if original:
          profile['address'] = address
          original = False
        else:
          new_profile['address'] = address
          profiles.append(new_profile)
          new_profile = profile.copy()
      else:
        print(f'Invalid DNS IP address:{address}. Fix and try again...')
        exit()
  
  return profiles

def build_session_acl_objects(profiles):
  """ Build session ACL objects from the TABLE_COLUMN dictionary. """

  aces_list = TABLE_COLUMNS['acl_sess.acl_sess__v4policy']
  for aces,profile in zip(aces_list,profiles):
    profile ['acc_sess__v4policy']= []
    add_entries_to_session_acl(aces,profile['acl_sess__v4policy'])
  
  return profiles

def add_entries_to_session_acl(aces,acl):
  """ Add the ace to the ACL """
  
  for ace in aces:
    acl.append(translate_ace_to_api_values(ace))
  
def translate_ace_to_api_values(ace):
  """ Given a typical ACE i.e. user any alias internal-networks any deny, extract the information and translate it to API specific values. """
  
  ace_values = ace.split(' ')
  
  ace_object = {}

  ace_values = add_src_dst_values_to_ace_object(ace_values,ace_object,'s')
  ace_values = add_src_dst_values_to_ace_object(ace_values,ace_object,'d')
  ace_values = add_service_to_ace_object(ace_values,ace_object)
  ace_values = add_action_values_to_ace_object(ace_values,ace_object)
  ace_values = add_extended_action_values_to_ace_object(ace_values,ace_object)

  return ace_object

def add_src_dst_values_to_ace_object(ace_values,ace_object,src_dst):
  """ Given a list of ACE values, extract the source information, add them to the ace_object and return a truncated list of ACE values. """
  allowed_values = set(['host','user','alias','network','role','localip','any'])
  src_value = ace_values[0]
  alias_prefix = 'src' if src_dst == 's' else 'dst'
  values_to_remove = 1

  if src_value not in allowed_values:
    print('Invalid entry for source value in ACE. Fix and try again.')
    exit()
  elif src_value == 'host':
    values_to_remove += 1
    if is_valid_ip_address(ace_values[values_to_remove-1]):
      ace_object[f'{src_dst}ipaddr'] = ace_values[values_to_remove-1]
      ace_object[alias_prefix] = f'{src_dst}host'
    else:
      print("Invalid host IP address for source value in ACE. Fix and try again.")
      exit()
  elif src_value == 'user':
    ace_object[alias_prefix] = f'{src_dst}user'
    ace_object[f'{src_dst}user'] = True
  elif src_value == 'any':
    ace_object[alias_prefix] = f'{src_dst}any'
    ace_object[f'{src_dst}any'] = True
  elif src_value == 'alias':
    ace_object[alias_prefix] = f'{src_dst}alias'
    ace_object[f'{alias_prefix}alias'] = ace_values[values_to_remove]
    values_to_remove += 1
  elif src_value == 'role':
    ace_object[alias_prefix] = f'{src_dst}userrole'
    ace_object[f'{src_dst}urname'] = ace_values[values_to_remove]
    values_to_remove += 1
  elif src_value == 'localip':
    ace_object[alias_prefix] = f'{src_dst}localip'
    ace_object[f'{src_dst}localip'] = True
  else:
    if len(ace_values[values_to_remove].split('.')) == 4:
      ipaddr,netmask = ace_values[values_to_remove].split('/')
      if is_valid_ip_address(ipaddr):
        ace_object[alias_prefix] = f'{src_dst}network'
        ace_object[f'{src_dst}network'] = ipaddr 
        ace_object[f'{src_dst}netmask'] = convert_subnetmask(netmask)
        values_to_remove += 1
      else:
        print('IP and or netmask not valid for source value in ACE. Aborting ...')
        exit()
  
  return ace_values[values_to_remove:]

def is_valid_ip_address(ipaddress):
  """ An IP address should be four integers separated by periods. """
  octets = ipaddress.split('.')
  for octet in octets:
    if octet.isdigit():
      octet_value = int(octet)
      if octet_value < 0 or octet_value > 255:
        return False
    else:
      return False
  
  return True

def convert_subnetmask(ipsubmask):
  """ Given a number between 1 and 32, convert to the decimal mask value. """

  octet_values = ['128','192','224','240','248','252','254','255']

  if ipsubmask.isdigit():
    subnet_number = int(ipsubmask)
    if subnet_number > 1 and subnet_number <= 32:
      iterations = math.ceil(subnet_number/8)
      netmask_string = ''
      while subnet_number > 0:
        if subnet_number <= 8:
          if netmask_string == '':
            netmask_string = f'{octet_values[subnet_number-1]}'
          else:
            netmask_string += f'.{octet_values[subnet_number-1]}'
          subnet_number -= subnet_number
        else:
          subnet_number -= 8
          if netmask_string == '':
            netmask_string = f'255'
          else:
            netmask_string += '.255'
      number_of_octets_added = len(netmask_string.split('.'))
      if number_of_octets_added < 4:
        iterations = 4 - number_of_octets_added
        while iterations > 0:
          netmask_string += '.0'
          iterations -= 1
      return netmask_string
    else:
      print("Subnet mask must be between 1 and 32. Fix and try again.")
      exit()
  else:
    print("Invalid value for subnet mask. Must be a number between 1 and 32. Fix and try again.")

def add_action_values_to_ace_object(ace_values,ace_object):
  """ Adds the action values to the ace object and returns the ace_values without the consumed action values. """

  action_value = ace_values[0]
  values_to_remove = 1

  if action_value == 'permit' or action_value == 'deny':
    ace_object[action_value] = True
    ace_object['action'] = action_value
  elif action_value == 'dst-nat':
    ace_object['action'] = action_value
    if ace_values[values_to_remove] == 'ip':
      ipaddr = ace_values[values_to_remove+1]
      if is_valid_ip_address(ipaddr):
        ace_object['dnatip'] = ace_values[values_to_remove+1]
        values_to_remove += 2
      else:
        print(f'Not a valid destination NAT IP:{ipaddr}. Fix and try again.')
        exit()
      if values_to_remove < len(ace_values):
        if ace_values[values_to_remove].isdigit():
          port_number = int(ace_values[values_to_remove])
          if port_number > 0 and port_number <= 65535:
            ace_object['dnatport'] = port_number
            values_to_remove += 1
          else:
            print('Port number must be between 0 and 65535. Fix and try again.')
            exit()
      else:
        return []
    elif ace_values[values_to_remove] == 'name':
      ace_object['dnathostname'] = ace_values[values_to_remove+1]
      values_to_remove += 2
    else:
      print("Destination NAT must specify an ip or name. Fix and try again.")
      exit()
  elif action_value == 'src-nat':
    ace_object['action'] = action_value
    ace_object['src-nat'] = True
    if values_to_remove < len(ace_values):
      if ace_values[values_to_remove] == 'pool':
        values_to_remove += 1
        ace_object['poolname'] = ace_values[values_to_remove]
  elif action_value == 'redirect':
    ace_object['action'] = 'redir_opt'
    redir_opt = ace_values[values_to_remove]
    values_to_remove += 1
    if redir_opt == 'tunnel':
      tunnel_id = ace_values[values_to_remove]
      if tunnel_id.isdigit():
        tunnel_id = int(tunnel_id)
        if tunnel_id >= 1 and tunnel_id <= 500:
          ace_object['tunid'] = tunnel_id
          ace_object['redir_opt'] = redir_opt
          values_to_remove += 1
      else:
        print('Redirect tunnel ID must be a number. Fix and try again.')
        exit()
    elif redir_opt == 'tunnel-group':
      ace_object['redir_opt'] = redir_opt
      ace_object['tungrpname'] = ace_values[values_to_remove]
      values_to_remove += 1
    elif redir_opt == 'group':
      ace_object['redir_opt'] = 'esi-group'
      ace_object['group'] = ace_values[values_to_remove]
      values_to_remove += 1
      if values_to_remove < len(ace_values):
        poss_direction = ace_values[values_to_remove]
        if poss_direction == 'direction':
          values_to_remove += 1
          if values_to_remove < len(ace_values):
            direction = ace_values[values_to_remove]
            if direction in ["forward","reverse","both"]:
              ace_object["dir"] = direction
              values_to_remove += 1
            else:
              print('Invalid value for direction: use forward, reverse or both. Fix and try again.')
              exit()
          else:
            print("Must specify forward, reverse or both after direction. Fix and try again.")
            exit()
  elif action_value == 'route':
    nat_type = ace_values[values_to_remove]
    if nat_type != 'src-nat' and nat_type != 'dst-nat':
      print('Route action must be specified with either dst-nat + ip/name or src-nat.')
      exit()
    values_to_remove += 1
    ace_object['action'] = 'route'
    if nat_type == 'dst-nat':
      try:
        name_or_ip = ace_values[values_to_remove]
        values_to_remove += 1
        value = ace_values[values_to_remove]
        if name_or_ip == 'name':
          ace_object['routednathostname'] = value
        elif name_or_ip == 'ip' and is_valid_ip_address(value):
          ace_object['routednatip'] = value
        else:
          raise ValueError
      except:
        print('route dst-nat must be followed by ip x.x.x.x or name hostname.')
        exit()
    nat_type_name = 'dst' if nat_type == 'dst-nat' else 'src'
    ace_object[f'{nat_type_name}-nat-route'] = True
  else:
    print("Invalid action. Specify permit, deny, redirect, src-nat, dst-nat, route src-nat or route dst-nat.")
      
  if values_to_remove == len(ace_values):
    return []
  else:
    return ace_values[values_to_remove:]    

def add_service_to_ace_object(ace_values,ace_object):
  """ Adds the service or protocol values to the ace_object. """

  ace_object['service_app'] = 'service'
  values_to_remove = 1
  proto_or_service = ace_values[0]
  
  if proto_or_service == 'udp' or proto_or_service == 'tcp':
    ace_object['svc'] = 'tcp_udp'
    ace_object['proto'] = proto_or_service
    proto = ace_values[values_to_remove]
    if proto.isdigit():
      proto = int(proto)
      if proto >= 1 and proto <= 65535:
        ace_object['port'] = 'range'
        ace_object['port1'] = proto
        values_to_remove += 1
        poss_port_range = ace_values[values_to_remove]
        if poss_port_range.isdigit():
          poss_port_range = int(poss_port_range)
          if poss_port_range >= 1 and poss_port_range <= 65535:
            if poss_port_range > proto:
              ace_object['port2'] = poss_port_range
              values_to_remove += 1
            else:
              print('Second port value must be smaller than first port value. Fix and try again.')
              exit()
          else:
            print("Ports must be from 1 to 255. Fix and try again.")
            exit()
        else:
          print("Port must be from 1 to 255. Fix and try again.")
      else:
        print("Port must be a number from 1 to 255. Fix and try again.")
    else:
      print("Port number or range must be specified.")
  elif proto_or_service.isdigit():
    protocol = int(proto_or_service)
    if protocol >= 1 and protocol <= 255:
      ace_object['svc'] = 'protocol'
      ace_object['protocol'] = protocol
    else:
      print('IP protocol must be a number between 1 and 255. Fix and try again.')
      exit()
  elif proto_or_service == 'any':
    ace_object['service-any'] = True
    ace_object['svc'] = 'service-any'
  elif is_service_name(proto_or_service):
    ace_object['svc'] = 'service-name'
    #double check that service exists on the network.
    ace_object['service-name'] = ace_values[values_to_remove-1]
  elif proto_or_service == 'icmp':
    #fill in later.
    print('icmp.')
  else:
    print('Invalid service keyword. Choose from tcp, udp, protocol, icmp, service or any.')
    exit()
  
  return ace_values[values_to_remove:]

def is_service_name(service):
  """ Checks for a name that starts with svc- or sys-svc- """
  start_pattern = re.compile(r'^svc-')
  system_defined_start = re.compile(r'^sys-svc-')

  if start_pattern.match(service) is not None or system_defined_start.match(service) is not None:
    return True
  else:
    #check that service exists on the network and return false if it doesn't.
    return False

def add_extended_action_values_to_ace_object(ace_values,ace_object):
  """ Add the extended actions to the ace object. """
  allowed_extended_actions = set(['blacklist','priority-802.1p','disable-scanning','tos','time-range','queue','log','mirror'])
  
  ext_action = ace_values[0]
  current_index = 1

  if ext_action not in allowed_extended_actions:
    print("Invalid extended action specified. Allowed are blacklist, priority-802.1p, disable-scanning, tos, time-range, queue, log, mirror")
  else:
    if ext_action == 'time-range':
      ace_object['trname'] = ace_values[current_index]
      current_index += 1
      if current_index < len(ace_values):
        return add_extended_action_values_to_ace_object(ace_values[current_index:],ace_object)
    elif ext_action == 'queue':
      ace_object['queue'] = True
      option = ace_values[current_index]
      if option == 'high' or option == 'low':
        ace_object['queue-type'] = option
      else:
        print("Queue must be specified as high or low. Fix and try again.")
        exit()
      current_index += 1
      if current_index < len(ace_values):
        return add_extended_action_values_to_ace_object(ace_values[current_index:],ace_object)
    elif ext_action == 'tos':
      tos_value = ace_values[current_index]
      if tos_value.isdigit():
        tos = int(tos_value)
        if tos >= 0 and tos <= 63:
         ace_object['tosstr'] = tos
        else:
          print("TOS value must be between 0 and 63 inclusive. Fix and try again.")
          exit()
        current_index += 1
      else:
        print('TOS value must be a number between 0 and 63 inclusive. Fix and try again.')
        exit()
      if current_index < len(ace_values):
        return add_extended_action_values_to_ace_object(ace_values[current_index:],ace_object)
    elif ext_action == 'priority-802.1p':
      dot1p_value = ace_values[current_index]
      if dot1p_value.isdigit():
        dot1p = int(dot1p_value)
        if dot1p >= 0 and dot1p <= 7:
          ace_object['prio8021p'] = dot1p
          current_index += 1
        else:
          print("802.1p value must be between 0 and 7. Fix and try again.")
          exit()
        if current_index < len(ace_values):
          return add_extended_action_values_to_ace_object(ace_values[current_index:],ace_object)
      else:
        print("802.1p value must be a number between 0 and 7. Fix and try again.")
        exit()
    else:
      ace_object[ext_action] = True
      if current_index < len(ace_values):
        return add_extended_action_values_to_ace_object(ace_values[current_index:],ace_object)

def add_acls_to_role(full_attribute_name,profiles):
  """Special function for handling roles and their associated ACLs. """
  
  acl_lists = TABLE_COLUMNS[f'{full_attribute_name}.pname']

  for acl_list,profile in zip(acl_lists,profiles):
    acls = acl_list.split(', ')
    for acl in acls:
      acl_type,pname = acl.split(' ')
      profile['role__acl'].append({'pname':pname,'acl_type':acl_type}) 

def add_auth_servers_to_server_group(profiles):
  """ There are special options for adding the authentication servers to the group. Adding all for example requires
      a boolean object to be added to the array. """

  attribute_name = 'server_group_prof.auth_server.name'
  # if not all, just add array attributes to the auth_server_group.
  if attribute_name in TABLE_COLUMNS:
    server_names = TABLE_COLUMNS[attribute_name]
    for server_name,profile in zip(server_names,profiles):
      profile['auth_server'] = []
      servers = server_name.replace(' ','').split(',')
      if len(servers) == 1:
        if servers[0] == 'All':
          profile['auth_server'].append({'all':True})
      elif servers == '':
        continue
      else:
        for server in servers:
          if is_valid_string_or_number(attribute_name,server):
            profile['auth_server'].append({'name':server})
          else:
            print('In add_servers_to_server_group: Invalid server name')
            exit()

def add_role_derivation_rules_to_server_group(profiles):
  ''' Server derived roles need to be added as derivation rules in the server group profile. '''

  derived_roles = TABLE_COLUMNS['server_group_prof.server_derived_role']
  
  for derived_role,profile in zip(derived_roles,profiles):
    if derived_role == 'True':
      profile['derivation_rules_vlan_role'] = [{'target':'role','value-of':True,'rule':'value-of','attribute':'role'}]
  
  return profiles

def add_dot1x_profile_attributes():
  """ Certain opmodes require specific dot1x profiles attributes to be added in the back end. """

  enterprise = ['wpa3-aes-ccm-128','wpa3-aes-gcm-256','wpa2-aes']
  #wpax-enterprise requires an eap method to be configured. Default is eap-peap and eap-mschapv2 for the inner/outer methods.
  if 'ssid_prof.opmode' in TABLE_COLUMNS:
    if not 'dot1x_auth_profile.profile-name' in TABLE_COLUMNS:
      TABLE_COLUMNS['dot1x_auth_profile.profile-name'] = TABLE_COLUMNS['ssid_prof.profile-name'].copy()
    opmodes = TABLE_COLUMNS['ssid_prof.opmode']
    if 'dot1x_auth_profile.termination_eaptype' in TABLE_COLUMNS:
      dot1x_outer_eap_types = TABLE_COLUMNS['dot1x_auth_profile.termination_eaptype']
    else:
      dot1x_outer_eap_types = ['EAP-PEAP' for _ in TABLE_COLUMNS['dot1x_auth_profile.profile-name']]
    dot1x_inner_eap_types = []
    dot1x_termination_eap = []
    add_eap_columns = False
    count = 0
    
    for opmode,outer_eap_type in zip(opmodes,dot1x_outer_eap_types):
      if  data_structures.BOOLEAN_DICT[opmode] in enterprise:
        if outer_eap_type != 'EAP-TLS':
          dot1x_outer_eap_types[count] = 'EAP-PEAP'
          dot1x_termination_eap.append('True')
        dot1x_inner_eap_types.append('EAP-MSCHAPV2')
        add_eap_columns = True
      else:
          dot1x_outer_eap_types[count] = ''
          dot1x_inner_eap_types.append('')
          dot1x_termination_eap.append('')
      count += 1

    if add_eap_columns:
      TABLE_COLUMNS['dot1x_auth_profile.termination_eaptype'] = dot1x_outer_eap_types
      TABLE_COLUMNS['dot1x_auth_profile.termination_innereaptype'] = dot1x_inner_eap_types
      TABLE_COLUMNS['dot1x_auth_profile.termination_mode'] = dot1x_termination_eap

def add_vlan_name_association(profiles,ordered_list):
  """ Make the VLAN name to ID associations if the IDs exist and have the same node information. """
  
  if 'vlan_name.name' in TABLE_COLUMNS and 'vlan_id.id' in TABLE_COLUMNS:
    vlan_name_profile_index,vlan_id_profile_index = 0,0
    vlan_name_profile_index = get_profile_name_index(ordered_list,'vlan_name.name')
    vlan_id_profile_index = get_profile_name_index(ordered_list,'vlan_id.id')
    vlan_name_profiles = profiles[vlan_name_profile_index]
    vlan_id_profiles = profiles[vlan_id_profile_index]
    vlan_name_id_profiles = [{} for _ in vlan_name_profiles]
    for vlan_name_profile,vlan_id_profile,vlan_name_id_profile in zip(vlan_name_profiles,vlan_id_profiles,vlan_name_id_profiles):
      if vlan_name_profile['node'] == vlan_id_profile['node']:
        vlan_name_id_profile['node'] = vlan_name_profile['node']
        vlan_name_id_profile['name'] = vlan_name_profile['name']
        vlan_name_id_profile['vlan-ids'] = str(vlan_id_profile['id'])
    if vlan_name_profile_index < vlan_id_profile_index:
      insert_index = vlan_id_profile_index + 1
    else:
      insert_index = vlan_name_profile_index + 1
    ordered_list.insert(insert_index,['vlan_name_id.name'])
    profiles.insert(insert_index,vlan_name_id_profiles)

  return profiles

def get_profile_name_index(ordered_list,profile_name):
  """ Returns the index for the vlan_name profiles. """
  
  index = 0
  for list in ordered_list:
    if profile_name in list:
      return index
    else:
      index += 1
  
  return -1

def add_array_attribute_to_profiles(full_attribute_name,profiles):
  """ Adds an array attribute to the profile. """

  prof_name,attribute_name = full_attribute_name.split('.')
  
  for profile in profiles:
    profile[attribute_name] = []
  if attribute_name == 'role__acl':
    add_acls_to_role(full_attribute_name,profiles)

  else:
    required_properties = get_required_properties(full_attribute_name)
  
    for required_property in required_properties:
      required_property_path = prof_name + '.' + attribute_name + '.' + required_property
      property_type = get_attribute_type(required_property_path)
      if required_property_path in TABLE_COLUMNS:
        attributes = TABLE_COLUMNS[required_property_path]

        if property_type == 'string' or property_type == 'integer':
          for attribute,profile in zip(attributes,profiles):
            attribute_value = None
            possible_attribute_numbers = extract_numbers_from_attribute(attribute)
            if len(possible_attribute_numbers) != 0:
              for number in possible_attribute_numbers:
                if is_valid_string_or_number(required_property_path,int(number),type=property_type):
                  attribute_value = int(number)
                  profile[attribute_name].append({required_property:attribute_value})
                else:
                  raise ValueError(f'Invalid value. {required_property} is incorrectly configured.')
            elif property_type == 'string' and is_enumerated_property(required_property_path):
              try:
               attribute_value = data_structures.BOOLEAN_DICT[attribute]
              except KeyError:
                print(f'Entry not in the BOOLEAN_DICT: Add {full_attribute_name}.{required_property}')
                exit()
            else:
              attribute_list = attribute.replace(' ','').split(',')
              if len(attribute_list) > 1:
                for attribute_item in attribute_list:
                  attribute_type = get_attribute_type(required_property_path)
                  if is_valid_string_or_number(required_property_path,attribute_item,attribute_type):
                    if attribute_type == 'integer':
                      attribute_item = int(attribute_item)
                    profile[attribute_name].append({required_property:attribute_item})
                  else:
                    raise ValueError(f'Invalid value. {required_property} is incorrectly configured.')
              elif len(attribute_list) == 1:
                attribute_item = attribute_list[0]
                profile[attribute_name].append({required_property:attribute_item})

def remove_empty_objects_that_are_not_booleans(full_attribute_name,attributes,profiles):
  """ Removes any attributes that were left empty in the table. """

  attribute_type = ''
  names = full_attribute_name.split('.')
  prof_name = names[0]
  attribute_name = names[1]
  
  if full_attribute_name in SPECIAL_COLUMNS:
    attribute_type = 'object'
  elif len(names) == 3 and get_attribute_type(f'{prof_name}.{attribute_name}') == 'array':
    attribute_type = get_attribute_type(full_attribute_name)
  else:    
    attribute_type = get_attribute_type(f'{prof_name}.{attribute_name}')
  
  if attribute_type == 'object':
    for attribute,profile in zip(attributes,profiles):
      if attribute == '' and len(profile[attribute_name].keys()) == 0:
        profile.pop(attribute_name)
  

def add_object_attribute_to_profiles(full_attribute_name,attributes,profiles):
  """ Adds the object attributes to the profiles. """

  _,attribute_name = full_attribute_name.split('.')
  required_properties = get_required_properties(full_attribute_name)

  for profile in profiles:
    profile[attribute_name] = {}

  if len(required_properties) == 0:
    properties = get_attribute_properties(full_attribute_name)
    if len(properties) != 0:
      add_boolean_attribute_to_profiles(full_attribute_name,attributes,profiles)

  else:
    for required_property in required_properties:
      required_property_path = full_attribute_name + '.' + required_property
      if required_property_path in TABLE_COLUMNS:
        add_attributes_to_profiles(required_property_path,TABLE_COLUMNS[required_property_path],profiles)

def get_attribute_properties(full_attribute_name):
  """ Returns a list of properties for the attribute or an empty list. """

  if len(full_attribute_name.split('.')) < 3:
    full_attribute_name += '.'

  prof_name,attribute_name,property_name = full_attribute_name.split('.')

  for ref in API_REF:
    if prof_name in ref['definitions']:
      if property_name != '':
        try:
          properties = [property for property in ref['definitions'][prof_name]['properties'][attribute_name]['properties'][property_name]['properties'].keys()]
        except KeyError:
          return []
      else:
        try:
          properties = [property for property in ref['definitions'][prof_name]['properties'][attribute_name]['properties'].keys()]
        except KeyError:
          return []
      return properties

def add_boolean_attribute_to_profiles(full_attribute_name,attributes,profiles):
  """ Adds boolean attributes to the profiles. """

  if len(full_attribute_name.split('.')) < 3:
    full_attribute_name += '.'
  
  _,attribute_name,property_name = full_attribute_name.split('.')

  for attribute,profile in zip(attributes,profiles):
    attribute = attribute.replace(' ','')
    attribute_list = attribute.split(',')
    value = ''
    for item in attribute_list:
      if item in data_structures.BOOLEAN_DICT.keys():
        if item == 'True' or item == 'False':
          if property_name != '':
            profile[attribute_name][property_name] = data_structures.BOOLEAN_DICT[item]
            return
          else:
            profile[attribute_name] = data_structures.BOOLEAN_DICT[item]
            return
        else:
          value = data_structures.BOOLEAN_DICT[item]
      elif item.isdigit():
        value = item
      else:
        raise ValueError(f'Invalid value for item {item} in {_}')

      if property_name != '':
        profile[attribute_name][property_name] = value
      else:
        profile[attribute_name][value] = True

def attribute_is_valid(full_attribute_name,attribute):
  """ Given a full attribute name, check the API to make sure the attribute is valid or not. """

  if len(full_attribute_name.split('.')) < 3:
    full_attribute_name += '.'
  prof_name,attribute_name,property_name = full_attribute_name.split('.')
  
  attribute_type = get_attribute_type(full_attribute_name)

  return is_valid_string_or_number(full_attribute_name,attribute,type=attribute_type)

def add_boolean_attributes_to_profile(prof_name,attribute_name,attribute,profile):
  """ Adds the boolean values to the attribute object in the profile. """

  boolean_list = attribute.split(',')
  digit_pattern = re.compile(r'\d+')

  for boolean in boolean_list:
    if digit_pattern.match(boolean) is not None:
      if property_is_in_properties(prof_name,attribute_name,boolean):
        profile[attribute_name][boolean] = True
    elif boolean in data_structures.BOOLEAN_DICT.keys():
      profile[attribute_name][boolean] = True
    else:
      print(f'Invalid value encountered in {prof_name}: {boolean} not an excepted value.')
      exit()

def property_is_in_properties(prof_name,attribute_name,property):
  """ Checks whether the property is in the properties of the given attribute. Returns True if so. """

  for ref in API_REF:
    if prof_name in ref["definitions"].keys():
      if property in ref["definitions"][prof_name]["properties"][attribute_name]["properties"]:
        return True
      else:
        return False

def get_required_properties(full_attribute_name):
  """ Returns the required properties list or an empty list. """

  prof_name,attribute_name= full_attribute_name.split('.')

  for ref in API_REF:
    if prof_name in ref["definitions"].keys():
      try:
        if get_attribute_type(full_attribute_name) == 'array':
          required = ref["definitions"][prof_name]["properties"][attribute_name]["items"]["required"]
        else:
          required = ref["definitions"][prof_name]["properties"][attribute_name]["required"]
        return required
      except KeyError:
        return []


def add_string_attribute(profile_name,attribute_name,attributes,profiles):
  """ Adds non-nested string attributes to the profiles. """

  for attribute,profile in zip(attributes,profiles):
    if is_valid_string_or_number(f'{profile_name}.{attribute_name}',attribute):
      profile[attribute_name] = attribute
    else:
      exit()

def is_valid_string_or_number(full_attribute_name,attribute,type="string"):
  """ Checks the API reference to ensure that the input is valid. """
  
  min_len = get_attribute_min_len(full_attribute_name)
  max_len = get_attribute_max_len(full_attribute_name)

  if type == "string":
    if is_enumerated_property(full_attribute_name):
      return string_is_in_enumerated_property_list(full_attribute_name,attribute)
    else:
      current_len = len(attribute)
  elif type == 'integer':
    current_len = int(attribute)
  else:
    print(f"in is_valid_string_or_number: type {type} is not supported. Choose from string or integer.")
    exit()
  
  return current_len >= min_len and current_len <= max_len

def get_attribute_min_len(full_attribute_name):
  """ Returns the minimal value for an attribute, property_name should exist for nested attributes. """

  names = full_attribute_name.split('.')
  property_name = ''
  if len(names) == 3:
    profile_name,attribute_name,property_name = names
  else:
    profile_name,attribute_name = names

  try:
    if property_name != "":
      for ref in API_REF:
        if profile_name in ref["definitions"]:
          if get_attribute_type(profile_name+'.'+attribute_name) == 'array':
            min_len = ref["definitions"][profile_name]["properties"][attribute_name]["items"]["properties"][property_name]["minimum"]
          else:
            min_len = ref["definitions"][profile_name]["properties"][attribute_name]["properties"][property_name]["minimum"]
    else:
      for ref in API_REF:
        if profile_name in ref["definitions"]:
          min_len = ref["definitions"][profile_name]["properties"][attribute_name]["minimum"]
  except KeyError:
    min_len = 1 if get_attribute_type(full_attribute_name) == 'string' else 0

  return min_len

def get_attribute_max_len(full_attribute_name):
  """ Returns the maximum value for an attribute, property_name should exist for nested attributes. """

  names = full_attribute_name.split('.')
  property_name = ''
  if len(names) == 3:
    profile_name,attribute_name,property_name = names
  else:
    profile_name,attribute_name = names

  try:
    if property_name != "":
      for ref in API_REF:
        if profile_name in ref["definitions"]:
          if get_attribute_type(profile_name+'.'+attribute_name) == 'array':
            max_len = ref["definitions"][profile_name]["properties"][attribute_name]["items"]["properties"][property_name]["maximum"]
          else:
            max_len = ref["definitions"][profile_name]["properties"][attribute_name]["properties"][property_name]["maximum"]
    else:
      for ref in API_REF:
        if profile_name in ref["definitions"]:
          max_len = ref["definitions"][profile_name]["properties"][attribute_name]["maximum"]
  except KeyError:
    max_len = 256 if get_attribute_type(full_attribute_name) == 'string' else 10000 

  return max_len

def attribute_has_properties(full_attribute_name):
  """ Checks whether the attribute has a properties key. Returns True if it does. """

  profile_name,attribute_name = full_attribute_name.split('.')

  for ref in API_REF:
    if profile_name in ref["definitions"]:
      if "properties" in ref["definitions"][profile_name][attribute_name].keys():
        return True
      else:
        return False

def get_object_property_name(full_attribute_name):
  """ Gets the property name of the object. """

  profile_name,attribute_name = full_attribute_name.split('.')

  for ref in API_REF:
    if profile_name in ref["definitions"]:
      return list(ref["definitions"][profile_name]["properties"][attribute_name]["properties"].keys())
  
  return []

def is_enumerated_property(full_attribute_name):
  """ Returns true if the property is an enumerated property. """
  names = full_attribute_name.split('.')
  property_name = ''
  if len(names) == 3:
    property_name = names[2]
  profile_name,attribute_name = names[0],names[1]

  for ref in API_REF:
    if profile_name in ref["definitions"]:
      if property_name != '':
        if get_attribute_type(f'{profile_name}.{attribute_name}') == 'array':
          return 'enum' in ref['definitions'][profile_name]["properties"][attribute_name]['items']['properties'][property_name]
        else:
          return "enum" in ref["definitions"][profile_name]["properties"][attribute_name]["properties"][property_name].keys()
      else:
        return 'enum' in ref['definitions'][profile_name]['properties'][attribute_name]

def string_is_in_enumerated_property_list(full_attribute_name,string):
  """ Validates whether the string is an allowed value for an enumerated object. """

  names = full_attribute_name.split('.')
  property_name = ''
  if len(names) == 3:
    property_name = names[2]
  profile_name,attribute_name = names[0],names[1]
  for ref in API_REF:
    if profile_name in ref["definitions"]:
      if property_name != '':
        if get_attribute_type(f'{profile_name}.{attribute_name}') == 'array':
          return string in ref['definitions'][profile_name]['properties'][attribute_name]['items']['properties'][property_name]['enum']
        return string in ref["definitions"][profile_name]["properties"][attribute_name]["properties"][property_name]["enum"]
      else:
        return string in ref['definitions'][profile_name]['properties'][attribute_name]['enum']

def get_attribute_type(full_attribute_name):
  """ Returns the type of the attribute as checked against the API_REF. """

  if len(full_attribute_name.split('.')) < 3:
    full_attribute_name += '.'
  profile_name,attribute_name,property_name = full_attribute_name.split('.')
  if property_name != '':
    is_array = get_attribute_type(f'{profile_name}.{attribute_name}') == 'array'

  for ref in API_REF:
    if profile_name in ref['definitions']:
      type = ''
      if property_name != '':
        if is_array:
          type = ref['definitions'][profile_name]['properties'][attribute_name]['items']['properties'][property_name]['type']
        else:
          type = ref['definitions'][profile_name]['properties'][attribute_name]['properties'][property_name]['type'] 
      else: 
        type = ref['definitions'][profile_name]['properties'][attribute_name]['type']
      return type

def extract_numbers_from_attribute(attribute):
  """ For integer attributes, remove any text that might have been included i.e. units like dBm. """
  
  has_letters = re.compile(r'[^\d]+')
  numbers = re.compile(r'\d+')
  has_units = re.compile(r'\d+ [a-zA-Z]+')
  extracted_numbers = numbers.findall(attribute)

  if len(extracted_numbers) != 0:
    if len(has_letters.findall(attribute)) != 0:
      if len(has_units.findall(attribute)) != 0:
        return extracted_numbers
      else:
        return []
    else:
      return extracted_numbers
  else:
    return extracted_numbers

# Device configuration specifics
def inventory_devices_in_network():
  """ Perform a query of the hierarchy and search for devices in the network. Return a dictionary of device name to MAC. """

  response = api.get_hierarchy()
  
  if response.status_code == 200:
    hierarchy = response.json()
    populate_devices_dict(hierarchy, DEVICE_DICTIONARY,'')
    return DEVICE_DICTIONARY
  else:
    return {}

def populate_devices_dict(hierarchy, devices_dict,current_node):
  """ Recursively build a device to MAC dictionary from the hierarchy given. """

  if hierarchy['name'] != '/':
    current_node += f'/{hierarchy["name"]}'
  if hierarchy['device_count'] > 0 and len(hierarchy['devices']) == 0:
    for childnode in hierarchy['childnodes']:
      populate_devices_dict(childnode,devices_dict,current_node)
  elif hierarchy['device_count'] > 0 and len(hierarchy['devices']) > 0:
    for device in hierarchy['devices']:
      devices_dict[device['name']] = f'{current_node}/{device["mac"]}'
  else:
    return 

def get_column_errors():
  """ Goes through the columns in the TABLE_COLUMNS and returns a list of errors that need to be fixed before the configuration objects
      will be built. """

  errors = []
  for table_column in TABLE_COLUMNS:
    if table_column in SPECIAL_COLUMNS:
      continue
    elif get_attribute_type(table_column) != 'boolean': 
      column_title = TABLE_COLUMNS[table_column][0]
      current_column = {column_title:[]}
      names = table_column.split('.')
      prof_name = names[0]
      for row,possible_list in enumerate(TABLE_COLUMNS[table_column][1:]):
        data = possible_list.replace(' ','').split(',')
        if table_column == OBJECT_IDENTIFIERS[prof_name]:
          data = [datum.split('%')[0] for datum in data]
        
        if get_attribute_type(table_column) == 'object':
          for col,datum in enumerate(data):
            if not is_valid_object(datum, table_column):
              current_column[column_title].append([row,col,datum])
        
        else:
          for col,datum in enumerate(data):
            if datum.isdigit():
              type = 'integer'
            else:
              type = 'string'
              if datum in data_structures.BOOLEAN_DICT and datum != 'True' and datum != 'False':
                datum = data_structures.BOOLEAN_DICT[datum]
            if datum == '':
              continue
            elif not is_valid_string_or_number(table_column,datum,type=type):
              current_column[column_title].append([row,col,datum])
        if len(current_column[column_title]) != 0:
          errors.append(current_column)
  
  return errors

def validate_COL_TO_ATTR_dict():
  """ Checks whether the COL_TO_ATTR dictionary contains endpoints in the API or not. """

  incorrect_api_path = {}

  for column_title in data_structures.COL_TO_ATTR:
    for assoc in data_structures.COL_TO_ATTR[column_title]:
      try:
        get_attribute_type(assoc)
      except KeyError:
        if column_title not in SPECIAL_COLUMNS:
          incorrect_api_path[column_title] = data_structures.COL_TO_ATTR[column_title]

  return incorrect_api_path

def remove_column_headers_from_columns_table():
  """ Column headers are required for the error checking in the get_column_errors function. Remove them for the rest of the program. """

  for column in TABLE_COLUMNS:
    TABLE_COLUMNS[column] = TABLE_COLUMNS[column][1:]

def is_valid_object(attribute,full_attribute_name):
  """ For objects for which the full attribute name is not three deep because the objects are booleans. We need to check the object list 
      similarly to enumerated types. """

  prof_name,attribute_name = full_attribute_name.split('.')
  for ref in API_REF:
    if prof_name in ref['definitions']:
      if not 'properties' in ref['definitions'][prof_name]['properties'][attribute_name]:
        result = True
      else:
        if attribute in data_structures.BOOLEAN_DICT:
          attribute = data_structures.BOOLEAN_DICT[attribute]
        result = attribute in ref['definitions'][prof_name]['properties'][attribute_name]['properties']
      return result

def is_nested_profile(profile_attribute):
  """ If an attribute points to another profile then it is a nested profile. Returns True or False. """

  nested_ending_abbr = re.compile(r'.+_prof$')
  nested_ending_full = re.compile(r'.+_profile$')

  if nested_ending_abbr.match(profile_attribute) is not None or nested_ending_full.match(profile_attribute) is not None:
    return True 

  else:
    return False 


def get_profile_names(profile_name_col, suffix=""):
  """ Gets the profile names column from the provided table and return a list of names appended with the optional suffix. """

  return [name.text+suffix for name in profile_name_col]

def create_table(columns):
    """ Returns a table created from the columns passed to it. """

    rows,cols = 0,0
    for _ in columns:
        cols += 1
    rows = len(columns[0])

    doc = Document()
    table = doc.add_table(cols=cols,rows=rows)
    for col,table_col in zip(columns,table.columns):
        for value,cell in zip(col,table_col.cells):
            cell.text = value
    
    return doc.tables

#external document handling
def get_table_titles(document):
    """ Given a word document name, returns an array of table_titles from the document. """
    design_doc = Document(document)

    paragraphs = design_doc.paragraphs
    title_match = re.compile(r'Table \d{1,2} . ((\w+\s?)+)')
    table_titles = []

    for paragraph in paragraphs:
        match = title_match.match(paragraph.text)
        if type(match) is re.Match:
            table_titles.append(match.group(1))

    return table_titles

def extract_row_cells_from_table(table):
  """ Returns an array of row_cells without the header row. """

  table_rows =[]

  enum_table_rows = enumerate(table.rows)

  for count,_ in enum_table_rows:
      table_rows.append(table.row_cells(count))

  return table_rows[1:]

def extract_column_cells_from_table(table):
  """ Returns an array of column_cells without the header column. """

  table_columns = []

  enum_table_cols = enumerate(table.columns)

  for count,_ in enum_table_cols:
    table_columns.append(table.column_cells(count))

  return table_columns[1:]

def get_unique_values_from_cells(table_cells):
    """ Returns a list of sites from the table column provided """
    values = []
    for cell in table_cells:
        if not cell.text in values:
            values.append(cell.text)
    
    return values

def get_column_from_table(column_name,table):
    """ Returns the column defined by column_name in the table. """
    title_row = table.row_cells(0)
    index = 0
    for title in title_row:
        if title.text == column_name:
            return table.column_cells(index)[1:]
        else:
            index += 1
    
    if index == len(table.columns):
        raise ValueError

def build_tables_columns_dict(tables):
  """ Build a dictionary of column_name : column_cells from the tables provided. """

  node_column = None

  for table in tables:
    table_columns = [column for column in table.columns]
    for column in table_columns:
      if column.cells[0].text == 'Node' or column.cells[0].text == 'Device':
        node_column = column.cells[1:]
        table_columns = remove_node_column(table)
        table_columns = add_node_info_to_profile_name(node_column, table_columns)
    
    for column in table_columns:
      if column.cells[0].text not in TABLE_COLUMNS.keys():
        try:
          attribute_names = data_structures.COL_TO_ATTR[column.cells[0].text]
          for attribute_name in attribute_names:
            TABLE_COLUMNS[attribute_name] = [sanitize_white_spaces(cell.text) for cell in column.cells]
        except KeyError:
          print(f"Your column {column.cells[0].text} is mistyped or the attribute is not currently supported. Delete the column or add the necessary information to the COL_TO_ATTR data structure in the data_structures.py file.")
          exit()

def remove_node_column(table):
  """ Removes the node column from the table. """

  table_columns = []
  
  for column in table.columns:
    if column.cells[0].text != 'Node' and column.cells[0].text != 'Device':
      table_columns.append(column)
  
  return table_columns

def add_entries_to_object_identifiers():
  """ Adds a list of entries to the objects identifiers set. """
  for ref in API_REF:
    for object in ref['definitions']:
      if 'required' in ref['definitions'][object]:
        if len(ref['definitions'][object]['required']) == 1:
          OBJECT_IDENTIFIERS[object] = f'{object}.{ref["definitions"][object]["required"][0]}'
        else:
          name_pattern = re.compile(r'.*name$')
          required_props = ref['definitions'][object]['required']
          for required_prop in required_props:
            if name_pattern.match(required_prop) is not None:
              if object not in OBJECT_IDENTIFIERS:
                OBJECT_IDENTIFIERS[object] = f'{object}.{required_prop}'

def add_node_info_to_profile_name(node_column, table_columns):
  """ Adds the node information to the profile names in the table. """

  for column in table_columns:
    node_info_added = False
    try:
      for attribute in data_structures.COL_TO_ATTR[column.cells[0].text]:
        attribute_name = attribute.split('.')[0]
        if attribute == OBJECT_IDENTIFIERS[attribute_name] and not node_info_added:
          current_node = node_column[0]
          for profile,node in zip(column.cells[1:],node_column):
            if node.text == '':
              node.text = current_node
            else:
              current_node = node.text
            if node.text in DEVICE_DICTIONARY:
              node_name = DEVICE_DICTIONARY[node.text]
            else:
              node_name = node.text
            profile.text += '%' + node_name
          node_info_added = True
    except KeyError:
      print(f"Your column {column.cells[0].text} is mistyped or the attribute is not currently supported. Delete the column or add the necessary information to the COL_TO_ATTR data structure in the data_structures.py file.")
  
  return table_columns
  
def build_profiles_dependencies(profiles_to_be_configured):
  """ Look for nested profiles and add entries to the TABLE_COLUMNS dictionary to configure those inner profiles. """

  copy = None

  for profile in profiles_to_be_configured:
    copy = profiles_to_be_configured.copy()
    copy.pop(profile)
    for other_profile in copy:
      if profile_is_an_attribute_of_current_profile(profile,other_profile):
        add_dependency_to_table_columns_dict(profile,other_profile)
        add_dependency_to_profiles_to_be_configured(profile,other_profile,profiles_to_be_configured)

def add_dependency_to_profiles_to_be_configured(current_profile,other_profile,profiles_to_be_configured):
  """ Adds the dynamic profiles to profiles to be configured. """

  other_prof_name = ''
  for property in get_profile_properties(current_profile):
    if property in data_structures.NESTED_DICT:
      if other_profile in data_structures.NESTED_DICT[property]:
        other_prof_name = property
    elif other_profile == property:
      other_prof_name = property

  other_profile_identifier = get_nested_object_identifier(f"{current_profile}.{other_prof_name}")

  other_prof_name += '.' + other_profile_identifier

  if other_prof_name not in profiles_to_be_configured[current_profile]:
    profiles_to_be_configured[current_profile].append(other_prof_name)  


def profile_is_an_attribute_of_current_profile(current_profile,other_profile):
  """ Returns True if other_profile is an attribute of current_profile. """
  
  current_profile_properties = get_profile_properties(current_profile)

  for property in current_profile_properties:
    if property in data_structures.NESTED_DICT:
      return other_profile in data_structures.NESTED_DICT[property]
    elif other_profile == property:
      return True
  return False

def add_dependency_to_table_columns_dict(current_profile,other_profile):
  """ Adds entries to the TABLE_COLUMNS dictionary if the two profiles are dependent. """

  current_profile_identifier = OBJECT_IDENTIFIERS[current_profile]
  other_profile_identifier = OBJECT_IDENTIFIERS[other_profile]
  if current_profile != '' and other_profile_identifier != '':
    current_profile_names = TABLE_COLUMNS[current_profile_identifier]
    other_profile_names = TABLE_COLUMNS[other_profile_identifier]
    dynamic_profile_name = get_dynamic_profile_name(current_profile, other_profile) 
    if dynamic_profile_name not in TABLE_COLUMNS:
      for profile_name in current_profile_names:
        if profile_name in other_profile_names:
          name_without_node_info = profile_name.split('%')[0]
          if dynamic_profile_name in TABLE_COLUMNS.keys():
            TABLE_COLUMNS[dynamic_profile_name].append(f'{name_without_node_info}_{other_profile}')
          else:
            TABLE_COLUMNS[dynamic_profile_name] = [f'{name_without_node_info}_{other_profile}'] 
        else:
          if dynamic_profile_name in TABLE_COLUMNS.keys():
            TABLE_COLUMNS[dynamic_profile_name].append('')
          else:
            TABLE_COLUMNS[dynamic_profile_name] = ['']

def get_dynamic_profile_name(current_profile,other_profile):
  """ Returns the attribute to be added between the two dependent profiles. """

  dynamic_profile_name = ''
  
  for property in get_profile_properties(current_profile):
    if property in data_structures.NESTED_DICT:
      if other_profile in data_structures.NESTED_DICT[property]:
        other_profile_identifier = get_nested_object_identifier(f"{current_profile}.{property}")
        dynamic_profile_name = f"{current_profile}.{property}.{other_profile_identifier}"
        return dynamic_profile_name
    elif property == other_profile:
      other_profile_identifier = get_nested_object_identifier(f"{current_profile}.{property}")
      dynamic_profile_name = f"{current_profile}.{property}.{other_profile_identifier}"
      return dynamic_profile_name

def get_nested_object_identifier(full_attribute_name):
  """ The identifier for the API object is not necessarily the same identifier for the API object nested inside another object. """

  prof_name,attribute_name = full_attribute_name.split('.')

  for ref in API_REF:
    if prof_name in ref['definitions']:
      if get_attribute_type(full_attribute_name) == 'array':
        inner_id = ref['definitions'][prof_name]['properties'][attribute_name]['items']['required']
      elif get_attribute_type(full_attribute_name) == 'object':
        inner_id = ref['definitions'][prof_name]['properties'][attribute_name]['required']
      else:
        inner_id = attribute_name
  
  if len(inner_id) > 1:
    has_name = re.compile(r'.*name$')
    for id in inner_id:
      if has_name.match(id) is not None:
        return id
  return inner_id[0]

def get_profile_properties(profile):
  """ Returns properties of the profile from the API_REF dictionary. """

  for ref in API_REF:
    if profile in ref['definitions']:
      return [prop for prop in ref['definitions'][profile]['properties']]

def get_profiles_to_be_configured():
  """ Get the profiles to be configured from the table keys. The result is a dictionary of profile: [attributes]. """

  profiles_to_configure = {}
  full_name = ''

  for name in TABLE_COLUMNS:
    if len(name.split('.')) < 3:
      name += '.'
    profile_name,attribute,property = name.split('.')
    if property != '':
      full_name = attribute + '.' + property
    else:
      full_name = attribute

    if profile_name in profiles_to_configure.keys():
      profiles_to_configure[profile_name].append(full_name)
    else:
      profiles_to_configure[profile_name] = [full_name]
  
  return profiles_to_configure

def check_that_required_attributes_are_provided(profile_name,profile_attributes,api_json_files):
  """ The profile is a dictionary of prof_name:[attributes]. The profile will be looked up in the API JSON files and the required attributes
      will be checked against the provided attributes. Returns a Boolean. """


  for api_json_file in api_json_files:
    if profile_name in api_json_file["definitions"].keys():
      required_attributes = api_json_file["definitions"][profile_name]["required"]
      for required_attribute in required_attributes:
        if required_attribute not in profile_attributes:
          print(f"You are missing a required attribute for the {profile_name} profile: {required_attribute} must be provided.")
          return False

  return True

def get_object_identifiers():
  """ Crawl through the API and get the required properties for all the objects. Returns a set of possible object identifiers. """

  object_ids = set() 

  for ref in API_REF:
    for ref_object in ref['definitions']:
      if 'required' in ref['definitions'][ref_object]:
        for required_prop in ref['definitions'][ref_object]['required']:
          object_ids.add(required_prop)        
  
  return object_ids

def get_objects_without_required_attributes():
  """ Returns a set of objects without required properties from the API. """

  objects = set()

  for ref in API_REF:
    for ref_object in ref['definitions']:
      if 'required' not in ref['definitions'][ref_object]:
        objects.add(ref_object)
  
  return objects
   
def sanitize_white_spaces(text):
  """ Remove leading/trailing white spaces and replace any carriage returns with spaces. """

  trailing_ws = re.compile(r'(.+) $')
  leading_ws = re.compile(r' (.+)$')

  match_trailing_ws = trailing_ws.match(text)
  if match_trailing_ws is not None:
    text = match_trailing_ws.group(1)
  
  match_leading_ws = leading_ws.match(text)
  if match_leading_ws is not None:
    text = match_leading_ws.group(1)

  text = text.replace('\n',' ')

  return text

def add_mac_auth_info_to_tables_columns():
  """ MAC Auth is set to either True or False in the columns. Use the ESSID information to make default
      MAC authentication profile names and AAA MAC auth profile parameters."""
  
  if 'mac_auth_profile.profile-name' in TABLE_COLUMNS:
    profile_names = []
    mac_auth_profile_names = []
    mac_auth_column = TABLE_COLUMNS.pop('mac_auth_profile.profile-name')
    if 'aaa_prof.profile-name' in TABLE_COLUMNS:
      profile_names = TABLE_COLUMNS['aaa_prof.profile-name']
    elif 'ssid_prof.profile-name' in TABLE_COLUMNS:
      profile_names = TABLE_COLUMNS['ssid_prof.profile-name']
      TABLE_COLUMNS['aaa_prof.profile-name'] = profile_names.copy()
    else:
      print('Either SSID profile and/or AAA profiles must be present for MAC Auth to be configured.')
      exit()
    for truth_value,profile_name in zip(mac_auth_column,profile_names):
      truthy = truth_value.split('%')[0]
      if truthy == 'True':
        mac_auth_profile_names.append(profile_name)
      else:
        mac_auth_profile_names.append('')
    TABLE_COLUMNS['mac_auth_profile.profile-name'] = mac_auth_profile_names
    TABLE_COLUMNS['aaa_prof.mac_auth_profile.profile-name'] = [name.split('%')[0] for name in mac_auth_profile_names] 

def is_object_in_api(attribute):
  """ Returns true for attributes that are objects in the API i.e. is a key in the 'definitions' object. """

  attribute_object_name = attribute.split('.')[0]

  for ref in API_REF:
    if attribute_object_name in ref['definitions']:
      return True
  
  return False

def is_a_nested_object(attribute,outer_object):
  """ Returns True if the attribute is an object in the API and is a property of the outer_object. """

  properties = get_object_properties(outer_object.split('.')[0])

  inner_name = ''

  if attribute in data_structures.BOOLEAN_DICT:
    inner_name = data_structures.BOOLEAN_DICT[attribute]
  else:
    inner_name = attribute.split('.')[0]
  
  if inner_name in properties and is_object_in_api(attribute):
    return True
  else:
    return False

def attribute_is_api_object(attribute):
  """ Returns True if the attribute found inside the object is itself an object in the API. """

  is_api_object = False

  if attribute in data_structures.NESTED_DICT:
    is_api_object = True
  else:
    is_api_object = is_object_in_api(attribute)

  return is_api_object  

def get_object_properties(object):
  """" Returns all the properties of an API object. """

  for ref in API_REF:
    if object in ref['definitions']:
      return ref['definitions'][object]['properties']
  
  return []

SPECIAL_COLUMNS = {
                   'reg_domain_prof.channel_width.width':add_wide_5ghz_channels,
                   'role.role__acl.pname':add_acls_to_role_profiles,
                   'server_group_prof.auth_server.name':add_auth_servers_to_server_group,
                   'server_group_prof.server_derived_role':add_role_derivation_rules_to_server_group,
                   'acl_sess.acl_sess__v4policy.ace':build_session_acl_objects,
                   'ip_dhcp_pool_cfg.ip_dhcp_pool_cfg__dns.address1': add_addresses_to_dhcp_pool,
                   'ip_dhcp_pool_cfg.ip_dhcp_pool_cfg__def_rtr.address': add_addresses_to_dhcp_pool,
                   'ip_dhcp_pool_cfg.ip_dhcp_pool_cfg__lease.var1': add_lease_to_dhcp_pool,
                   'ip_dhcp_pool_cfg.ip_dhcp_pool_cfg__lease.var2': add_lease_to_dhcp_pool,
                   'ip_dhcp_pool_cfg.ip_dhcp_pool_cfg__lease.var3': add_lease_to_dhcp_pool,
                   'ip_name_server.address': add_dns_server_addresses,
                   }